# -*- coding: utf-8 -*-
"""Assistant widget API routes.

Supported providers:

- OpenAI ChatGPT (proxying to OpenAI's API)
- Custom webhooks
- OpenWebUI-compatible chat APIs
- Embedded MCP server tooling
- Built-in tooling that answers by querying the local PostgreSQL database
  across Tickets, Knowledge Base, Inventory, and Network modules.
"""

import json
import re
import os
import uuid
import mimetypes
import ipaddress
import time
from collections import Counter
from datetime import date, datetime, timedelta
from decimal import Decimal, InvalidOperation
from typing import List, Dict, Any, Iterable, Optional, Tuple, Union, Sequence, Set

import requests
from flask import Blueprint, jsonify, request, abort, current_app, url_for, send_from_directory
from flask_login import login_required, current_user
from flask_babel import gettext as _
from sqlalchemy import func, or_, and_, literal
from sqlalchemy.orm import joinedload
from app.utils.files import secure_filename

from app import csrf, db
from app.models import (
    AssistantConfig,
    AssistantSession,
    AssistantMessage,
    AssistantDocument,
    Ticket,
    TicketComment,
    Attachment as TicketAttachment,
    AuditLog,
    KnowledgeArticle,
    KnowledgeArticleVersion,
    KnowledgeAttachment,
    HardwareAsset,
    SoftwareAsset,
    Network,
    NetworkHost,
    Contract,
    AddressBookEntry,
    User,
    TapeCartridge,
    TapeLocation,
)
from app.models.assistant import DEFAULT_SYSTEM_PROMPT
from app import db
from app.navigation import is_feature_allowed


CIDR_PATTERN = re.compile(r"\b\d{1,3}(?:\.\d{1,3}){3}/\d{1,2}\b")
IP_ADDRESS_PATTERN = re.compile(r"\b\d{1,3}(?:\.\d{1,3}){3}\b")
HOSTNAME_QUERY_PATTERN = re.compile(r"(?:hostname|host)\s+(?:named\s+)?([a-z0-9_.\-\:]+)", re.IGNORECASE)
MAC_QUERY_PATTERN = re.compile(r"(?:mac(?:\s+address)?)\s+(?:of|for|=)?\s*([0-9a-f:\-]{8,})", re.IGNORECASE)
TICKET_ID_PATTERN = re.compile(r"ticket\s+#?(\d+)|αιτημ(?:α|ατος)?\s*#?(\d+)", re.IGNORECASE)
USER_REF_PATTERN = re.compile(
    r"(?:assigned to|for|owner|user|ip\s+of|για|σε)\s+([a-z0-9_.\-άέήίόύώ\s]+)",
    re.IGNORECASE,
)
CREATED_BY_PATTERN = re.compile(
    r"created by\s+([a-z0-9_.\-άέήίόύώ\s]+)",
    re.IGNORECASE,
)
ASSET_TAG_PATTERN = re.compile(r"(?:asset|ετικέτα)\s+tag\s+#?([a-z0-9_.\-]+)", re.IGNORECASE)
HOSTNAME_PATTERN = re.compile(r"(?:hostname|όνομα\s+host)\s+([a-z0-9_.\-]+)", re.IGNORECASE)
SERIAL_PATTERN = re.compile(r"(?:serial|σειριακό)\s+#?([a-z0-9_.\-]+)", re.IGNORECASE)
SOFTWARE_TAG_PATTERN = re.compile(r"(?:license|software|app|άδεια|λογισμικό)\s+tag\s+#?([a-z0-9_.\-]+)", re.IGNORECASE)
STATUS_PATTERN = re.compile(r"status\s*(?:is|=|:)?\s*([a-z0-9_\- ]+)", re.IGNORECASE)
PRIORITY_PATTERN = re.compile(r"priority\s*(?:is|=|:)?\s*([a-z0-9_\- ]+)", re.IGNORECASE)
CREATED_ON_PATTERN = re.compile(r"created\s+(?:on|at|=)\s+([0-9/\-]+)", re.IGNORECASE)
CREATED_AFTER_PATTERN = re.compile(r"created\s+(?:after|since)\s+([0-9/\-]+)", re.IGNORECASE)
CREATED_BEFORE_PATTERN = re.compile(r"created\s+(?:before|until|prior to)\s+([0-9/\-]+)", re.IGNORECASE)
CREATED_BETWEEN_PATTERN = re.compile(
    r"created\s+(?:between|from)\s+([0-9/\-]+)\s+(?:and|to)\s+([0-9/\-]+)",
    re.IGNORECASE,
)
CLOSED_ON_PATTERN = re.compile(r"closed\s+(?:on|at|=)\s+([0-9/\-]+)", re.IGNORECASE)
CLOSED_AFTER_PATTERN = re.compile(r"closed\s+(?:after|since)\s+([0-9/\-]+)", re.IGNORECASE)
CLOSED_BEFORE_PATTERN = re.compile(r"closed\s+(?:before|until|prior to)\s+([0-9/\-]+)", re.IGNORECASE)
CLOSED_BETWEEN_PATTERN = re.compile(
    r"closed\s+(?:between|from)\s+([0-9/\-]+)\s+(?:and|to)\s+([0-9/\-]+)",
    re.IGNORECASE,
)
GENERIC_BETWEEN_PATTERN = re.compile(
    r"(?:between|from)\s+([0-9/\-]+)\s+(?:and|to)\s+([0-9/\-]+)",
    re.IGNORECASE,
)
CONTRACT_NUMBER_PATTERN = re.compile(
    r"(?:contract(?:\s+(?:number|no\.?|#))?|συμβ(?:όλαιο|ολαιο|άση|αση)\s*(?:αριθ(?:μός|μος)?|#)?)\s*[:#]?\s*([a-z0-9\-_/]+)",
    re.IGNORECASE,
)
DATE_BY_PATTERN = re.compile(r"(?:by|before|until|έως|εως|μέχρι|μεχρι)\s+([0-9/\-]+)", re.IGNORECASE)
AMOUNT_THRESHOLD_PATTERN = re.compile(r"(?:over|greater than|above|άνω των|ανω των)\s*[$€]?\s*([0-9.,]+)", re.IGNORECASE)
VENDOR_PATTERN = re.compile(
    r"(?:vendor|προμηθευτ[ήςη])\s+(?:is|=|:|for|με|του|τον|την)?\s*['\"]?([a-z0-9_.\-άέήίόύώ\s]+)",
    re.IGNORECASE,
)
VENDOR_TRAILING_PATTERN = re.compile(
    r"\b([a-z0-9_.\-άέήίόύώ]+)\s+vendor\b",
    re.IGNORECASE,
)
CONTRACT_FROM_PATTERN = re.compile(
    r"contracts?\s+(?:from|by)\s+([a-z0-9_.\-άέήίόύώ\s]+)",
    re.IGNORECASE,
)
CONTRACT_FOR_PATTERN = re.compile(
    r"contracts?\s+(?:for|about|regarding)\s+([a-z0-9_.\-άέήίόύώ\s]+)",
    re.IGNORECASE,
)
CONTRACT_TYPE_PATTERN = re.compile(r"(?:type|τύπος|τυπος)\s+(?:is|=|:)?\s*['\"]?([a-z0-9_.\-άέήίόύώ\s]+)", re.IGNORECASE)
COMPANY_PATTERN = re.compile(r"(?:company|εταιρεία|εταιρεια)\s+(?:is|=|:|at|στο|στη|στην|στης)?\s*['\"]?([a-z0-9_.\-άέήίόύώ\s]+)", re.IGNORECASE)
DEPARTMENT_PATTERN = re.compile(r"(?:department|dept|τμήμα|τμημα)\s+(?:is|=|:|στο|στη|στην|στον|σε)?\s*['\"]?([a-z0-9_.\-άέήίόύώ\s]+)", re.IGNORECASE)
CITY_PATTERN = re.compile(r"(?:city|πόλη|πολη)\s+(?:is|=|:|στο|στη|στην)?\s*['\"]?([a-z0-9_.\-άέήίόύώ\s]+)", re.IGNORECASE)
TAG_PATTERN = re.compile(r"(?:tag|ετικέτα|ετικετα)\s+(?:is|=|:)?\s*['\"]?([a-z0-9_.\-άέήίόύώ\s]+)", re.IGNORECASE)
LOCATION_PATTERN = re.compile(r"(?:location|τοποθεσία|τοποθεσια)\s+(?:is|=|:|στο|στη|στην|στον|σε)?\s*['\"]?([a-z0-9_.\-άέήίόύώ\s]+)", re.IGNORECASE)
EMAIL_DOMAIN_PATTERN = re.compile(r"(?:@|domain\s+)([a-z0-9_.\-]+\.[a-z]{2,})", re.IGNORECASE)
PHONE_PATTERN = re.compile(r"(?:phone|τηλέφωνο|τηλεφωνο)\s*(?:is|=|:)?\s*([+0-9()\s\-]{6,})", re.IGNORECASE)
CONTACT_NAME_PATTERN = re.compile(
    r"(?:contact|επάφη|επαφή|επαφη)\s+(?:details\s+for|for|named|=|:)?\s*['\"]?([a-z0-9_.\-άέήίόύώ\s]+)",
    re.IGNORECASE,
)
SUBJECT_QUOTED_PATTERN = re.compile(r"subject\s*(?:is|=|contains|like|:)?\s*\"([^\"]+)\"", re.IGNORECASE)
SUBJECT_SINGLE_QUOTED_PATTERN = re.compile(r"subject\s*(?:is|=|contains|like|:)?\s*'([^']+)'", re.IGNORECASE)
DESCRIPTION_QUOTED_PATTERN = re.compile(r"description\s*(?:is|=|contains|like|:)?\s*\"([^\"]+)\"", re.IGNORECASE)
DESCRIPTION_SINGLE_QUOTED_PATTERN = re.compile(r"description\s*(?:is|=|contains|like|:)?\s*'([^']+)'", re.IGNORECASE)
SITE_FILTER_PATTERN = re.compile(
    r"(?:site|location)\s+(?:is|=|:|at|στο|στη|στην|στον|στης)?\s*['\"]?([a-z0-9_.\-άέήίόύώ\s]+)",
    re.IGNORECASE,
)

STOP_WORDS = {
    "the",
    "and",
    "for",
    "with",
    "that",
    "from",
    "this",
    "what",
    "all",
    "any",
    "where",
    "when",
    "which",
    "are",
    "was",
    "can",
    "you",
    "give",
    "show",
    "list",
    "tell",
    "about",
    "please",
    "help",
    "need",
    "information",
    "details",
    "lookup",
    "find",
    "me",
    "my",
    "latest",
    "assign",
    "assigned",
    "assignment",
    "assignments",
    "unassigned",
    "pdf",
    "file",
    "document",
}

OPEN_STATUS_VALUES = ["open", "in progress", "pending", "new", "reopened"]
CLOSED_STATUS_VALUES = ["closed", "resolved", "completed", "cancelled", "done"]
PRIORITY_LEVELS = ["critical", "urgent", "high", "medium", "low"]

USER_REFERENCE_DELIMITERS = (
    " and ",
    " with ",
    " without ",
    " where ",
    " whose ",
    " having ",
    " status ",
    " priority ",
    " department ",
    " created ",
    " updated ",
    " closed ",
    " open ",
    " reopened ",
    " resolved ",
    " pending ",
    " tickets",
    " ticket",
    " incidents",
    " incident",
    " requests",
    " request",
)

TICKET_KEYWORDS = {
    "ticket", "tickets", "incident", "case", "request",
    "αίτημα", "αιτήματα", "εισιτήριο", "εισιτήρια", "υπόθεση"
}
TICKET_TEXT_STOP = {
    "ticket", "tickets", "incident", "incidents", "case", "cases", "request", "requests",
    "open", "closed", "pending", "progress", "status", "priority", "department",
    "assigned", "assign", "assignee", "me", "my", "today", "yesterday", "count", "how",
    "many", "list", "show", "find", "latest", "new", "update", "updated"
}
KNOWLEDGE_KEYWORDS = {
    "knowledge", "kb", "article", "articles", "procedure", "guide", "manual", "attachment", "attachments",
    "info", "information", "details",
    "γνωση", "γνώση", "γνώσεων", "άρθρο", "άρθρα", "διαδικασία", "οδηγία", "εγχειρίδιο", "συνημμένο", "συνημμένα"
}
SOFTWARE_KEYWORDS = {
    "software", "license", "licence", "application", "app", "subscription", "saas", "programme", "key", "keys", "product key", "serial", "serial number",
    "λογισμικό", "άδεια", "άδειες", "εφαρμογή", "εφαρμογές", "συνδρομή", "κλειδί", "κλειδιά", "σειριακό", "σειριακός"
}

HARDWARE_KEYWORDS = {
    "hardware", "device", "devices", "laptop", "desktop", "server", "asset", "equipment", "machine", "serial", "serial number", "asset tag",
    "υλικό", "συσκευή", "συσκευές", "υπολογιστής", "φορητός", "σταθερός", "server", "περιουσιακό", "εξοπλισμός", "σειριακό", "ετικέτα", "asset"
}
NETWORK_KEYWORDS = {
    "network", "subnet", "cidr", "vlan", "ip", "wifi",
    "δίκτυο", "υποδίκτυο", "cidr", "vlan", "ip", "ίπ", "δικτυο",
    "networks", "subnets"
}

CONTRACT_KEYWORDS = {
    "contract", "contracts", "renewal", "renewals", "vendor", "support", "coverage", "auto-renew", "auto renew",
    "agreement", "po", "purchase order", "sla", "service level", "value", "amount",
    "συμβαση", "συμβάση", "συμβασεις", "συμβάσεις", "ανανεωση", "ανανεώσεις", "προμηθευτη", "προμηθευτή",
    "υποστήριξη", "υποστηριξη", "τυπος", "τύπος", "συμβολαιο", "συμβόλαιο"
}

CONTRACT_VENDOR_STOPWORDS = {
    "contract", "contracts", "vendor", "vendors", "support", "phone", "email", "contact",
    "for", "about", "regarding", "with", "and", "the", "a", "an", "show", "list", "find",
    "search", "lookup", "give", "me", "please", "all", "any", "value", "values", "cost",
    "costs", "price", "prices", "details", "info", "information", "help", "θέλω", "να", "βρες",
    "renewals", "renewal", "status", "statuses", "type", "types"
}

COST_KEYWORDS = {
    "cost", "costs", "price", "value", "amount", "total", "budget",
    "κόστος", "κοστος", "τιμή", "τιμη", "ποσο", "πόσο", "ποσο κοστίζει", "κοστίζει"
}

ADDRESS_BOOK_KEYWORDS = {
    "contact", "contacts", "address book", "directory", "phone", "mobile", "email", "company", "department", "city", "tag",
    "vendor", "partner", "stakeholder",
    "επαφη", "επαφή", "επαφες", "επαφές", "τηλεφωνο", "κινητο", "email", "εταιρεια", "εταιρεία",
    "τμημα", "τμήμα", "πολη", "πόλη", "ετικετα", "ετικέτα", "συνεργατες", "συνεργάτες", "προμηθευτες", "προμηθευτές"
}

BACKUP_KEYWORDS = {
    "backup", "tape", "tapes", "lto", "cartridge", "cartridges",
    "disk", "disks", "drive", "drives",
    "backup job", "backup jobs",
    "ταινια", "ταινία", "ταινιες", "ταινίες", "κασέτα", "κασέτες", "κασετα", "κασετες",
    "δίσκος", "δίσκοι", "δισκος", "δισκοι",
}

BACKUP_LOCATION_LABELS = {
    "on_site": "On-Site",
    "in_transit": "In Transit",
    "off_site": "Off-Site",
}

MESSAGE_ALIAS_MAP = {
    "σήμερα": "today",
    "σημερα": "today",
    "χτες": "yesterday",
    "χθες": "yesterday",
    "ενημερώθηκαν": "updated",
    "ενημερωθηκαν": "updated",
    "δημιουργήθηκαν": "created",
    "δημιουργηθηκαν": "created",
    "δημιουργημένα": "created",
    "δημιουργημενα": "created",
    "δημιουργημένα από": "created by",
    "δημιουργημενα απο": "created by",
    "έκλεισαν": "closed",
    "εκλεισαν": "closed",
    "ανοικτά": "open",
    "ανοιχτά": "open",
    "ανοικτα": "open",
    "ανοιχτα": "open",
    "κλειστά": "closed",
    "κλειστα": "closed",
    "υψηλής προτεραιότητας": "high priority",
    "υψηλης προτεραιοτητας": "high priority",
    "μη ανατεθειμένα": "unassigned",
    "μη ανατεθειμενα": "unassigned",
    "χωρίς ανάθεση": "unassigned",
    "χωρις αναθεση": "unassigned",
    "ανατεθειμένα στον": "assigned to",
    "ανατεθειμενα στον": "assigned to",
    "ανατεθειμένα στη": "assigned to",
    "ανατεθειμενα στη": "assigned to",
    "εκπρόθεσμα": "overdue",
    "εκπροθεσμα": "overdue",
    "συνημμένα": "attachments",
    "συνημμενα": "attachments",
    "σχόλια": "comments",
    "σχολια": "comments",
    "ιστορικό ενεργειών": "audit log",
    "ιστορικο ενεργειων": "audit log",
    "τμήμα": "department",
    "τμημα": "department",
    "τοποθεσία": "location",
    "τοποθεσια": "location",
    "κατάσταση": "status",
    "κατασταση": "status",
    "ανανεώσεις": "renewals",
    "ανανεωση": "renewals",
    "λήγουν": "ending",
    "ληγουν": "ending",
    "αυτόματη ανανέωση": "auto-renew",
    "αυτοματη ανανεωση": "auto-renew",
    "επανάνοιγμα": "reopen",
    "επαννοιγμα": "reopen",
    "δεσμευμένες": "reserved",
    "δεσμευμενες": "reserved",
    "χωρίς ανάθεση": "unassigned",
    "χωρις αναθεση": "unassigned",
    "ταινίες": "tapes",
    "ταινιες": "tapes",
    "ταινία": "tape",
    "ταινια": "tape",
    "κασέτες": "tapes",
    "κασετες": "tapes",
    "κασέτα": "tape",
    "κασετα": "tape",
    "δίσκος": "disk",
    "δίσκοι": "disks",
    "δισκος": "disk",
    "δισκοι": "disks",
    "επαφή": "contact",
    "επαφη": "contact",
    "επαφές": "contacts",
    "επαφες": "contacts",
    "εταιρεία": "company",
    "εταιρεια": "company",
    "πόλη": "city",
    "πολη": "city",
    "ετικέτα": "tag",
    "ετικετα": "tag",
    "συμβάσεις": "contracts",
    "συμβασεις": "contracts",
    "συμβόλαιο": "contract",
    "συμβολαιο": "contract",
    "υποστήριξη": "support",
    "υποστηριξη": "support",
    "προμηθευτή": "vendor",
    "προμηθευτη": "vendor",
    "διαδικασίες": "procedures",
    "διαδικασιες": "procedures",
    "άρθρα": "articles",
    "αρθρα": "articles",
    "άρθρο": "article",
    "αρθρο": "article",
    "δημοσιευμένων": "published",
    "δημοσιευμενων": "published",
    "πρόχειρων": "draft",
    "προχειρων": "draft",
    "εγγύηση": "warranty",
    "εγγυηση": "warranty",
    "λήγει": "expiring",
    "ληγει": "expiring",
    "εκτός εγγύησης": "out of warranty",
    "εκτος εγγυησης": "out of warranty",
    "αποσύρμενα": "decommissioned",
    "αποσυρμενα": "decommissioned",
    "εγκατάστασης": "deployment",
    "εγκαταστασης": "deployment",
    "δεσμευμένες ip": "reserved ips",
    "χωρίς ανάθεση hosts": "unassigned hosts",
    "μεταξύ": "between",
    "μεταξυ": "between",
    "έως": "to",
    "εως": "to",
    "ως": "to",
    "αναζήτηση": "search",
    "αναζητηση": "search",
    "εμφάνισε": "show",
    "εμφανισε": "show",
    "δείξε": "show",
    "δειξε": "show",
    "λίστα": "list",
    "λιστα": "list",
    "βρες": "find",
    "άδειες": "licenses",
    "αδειες": "licenses",
    "συνδρομητικές": "subscription",
    "συνδρομητικες": "subscription",
    "διαχρονικές": "perpetual",
    "διαχρονικες": "perpetual",
    "τελευταίες 7 μέρες": "last 7 days",
    "τελευταιες 7 μερες": "last 7 days",
    "τελευταίων 7 ημερών": "last 7 days",
    "τελευταιων 7 ημερων": "last 7 days",
    "ind": "find",
    "fing": "find",
    "serach": "search",
    "crontract": "contract",
    "crontracts": "contracts",
    "contracs": "contracts",
    "contratcs": "contracts",
}

NETWORK_HINT_STOPWORDS = {
    "map",
    "maps",
    "diagram",
    "diagrams",
    "topology",
    "topologies",
    "layout",
    "layouts",
    "overview",
    "summary",
    "details",
    "information",
    "info",
    "hosts",
    "host",
    "available",
    "addresses",
    "address",
    "ips",
    "ip",
    "list",
    "lists",
    "show",
    "find",
    "lookup",
    "questions",
    "question",
    "from",
    "database",
    "db",
    "inventory",
    "status",
    "update",
    "updates",
    "report",
    "reports",
    "what",
    "is",
    "the",
    "this",
    "that",
    "which",
    "who",
    "where",
    "gateway",
    "vlan",
    "site",
    "hosts",
    "host",
    "network",
    "subnet",
    "cidr",
}

NETWORK_HINT_BREAK_RE = re.compile(r"\b(for|with|from|of|to|in|on|about|regarding|covering|showing)\b", re.IGNORECASE)

SOFTWARE_SEARCH_COLUMNS = [
    SoftwareAsset.name,
    SoftwareAsset.vendor,
    SoftwareAsset.category,
    SoftwareAsset.version,
    SoftwareAsset.license_type,
    SoftwareAsset.license_key,
    SoftwareAsset.serial_number,
    SoftwareAsset.custom_tag,
    SoftwareAsset.platform,
    SoftwareAsset.environment,
    SoftwareAsset.status,
    SoftwareAsset.cost_center,
    SoftwareAsset.support_vendor,
    SoftwareAsset.support_email,
    SoftwareAsset.support_phone,
    SoftwareAsset.contract_url,
    SoftwareAsset.usage_scope,
    SoftwareAsset.deployment_notes,
]

SENSITIVE_LICENSE_TERMS = (
    "license key",
    "licence key",
    "product key",
    "serial number",
    "activation key",
    "activation code",
    "windows key",
    "software key",
    "cd key",
    "κλειδί",
    "σειριακό",
)

REFUSAL_MARKERS = (
    "unable to provide",
    "cannot provide",
    "can't provide",
    "not able to provide",
    "cannot help with that request",
    "can't help with that request",
)

BUILTIN_DEFAULT_RESPONSE = (
    "I can help with ticket summaries, knowledge base articles and attachments, inventory assignments, "
    "and network availability. For example: 'Show my open tickets today', 'Which articles mention VPN', "
    "'List hardware assigned to Alice', or 'Find an available IP in 192.168.1.0/24'."
)

LLM_TOOL_DEFINITIONS: List[Dict[str, Any]] = []

PHRASE_PATTERN = re.compile(
    r"(?:for|about|regarding|lookup|find|search for|αναζήτησε|βρες|για)\s+([a-z0-9\"'\\-\\s\\.]+)",
    re.IGNORECASE,
)
QUOTED_PATTERN = re.compile(r"\"([^\"]+)\"|'([^']+)'")

assistant_bp = Blueprint("assistant", __name__, url_prefix="/assistant")
csrf.exempt(assistant_bp)


def _load_config():
    config = AssistantConfig.get()
    if not config or not config.is_enabled:
        return None
    return config


def _assistant_upload_folder() -> str:
    upload_folder = current_app.config.get("ASSISTANT_UPLOAD_FOLDER")
    if not upload_folder:
        upload_folder = os.path.join(current_app.instance_path, "assistant_uploads")
        current_app.config["ASSISTANT_UPLOAD_FOLDER"] = upload_folder
    os.makedirs(upload_folder, exist_ok=True)
    return upload_folder


def _extract_document_text(file_path: str, mimetype: Optional[str]) -> Optional[str]:
    mimetype = (mimetype or "").lower()
    try:
        if mimetype.startswith("text/") or mimetype in {"application/json", "application/xml"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as handle:
                return handle.read()
        if mimetype == "application/pdf":
            try:
                from pdfminer.high_level import extract_text

                return extract_text(file_path)
            except Exception:
                return None
        if mimetype in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }:
            try:
                import docx

                document = docx.Document(file_path)
                return "\n".join(paragraph.text for paragraph in document.paragraphs)
            except Exception:
                return None
    except Exception:
        return None
    return None


def _serialize_message(message: AssistantMessage) -> Dict[str, Any]:
    return {
        "id": message.id,
        "session_id": message.session_id,
        "role": message.role,
        "content": message.content,
        "created_at": message.created_at.isoformat() if message.created_at else None,
    }


def _serialize_document(document: AssistantDocument) -> Dict[str, Any]:
    return {
        "id": document.id,
        "session_id": document.session_id,
        "filename": document.original_filename,
        "mimetype": document.mimetype,
        "size": document.file_size,
        "status": document.status,
        "created_at": document.created_at.isoformat() if document.created_at else None,
        "download_url": url_for(
            "assistant.download_document", session_id=document.session_id, document_id=document.id
        ),
    }


def _serialize_session(session: AssistantSession) -> Dict[str, Any]:
    return {
        "id": session.id,
        "title": session.title,
        "created_at": session.created_at.isoformat() if session.created_at else None,
        "updated_at": session.updated_at.isoformat() if session.updated_at else None,
        "messages": [
            _serialize_message(message)
            for message in sorted(
                session.messages,
                key=lambda m: (m.created_at or date.min, m.id),
            )
        ],
        "documents": [_serialize_document(doc) for doc in session.documents],
    }


def _ensure_session_for_user(session_id: int, user: User) -> Optional[AssistantSession]:
    if not session_id:
        return None
    return AssistantSession.query.filter_by(id=session_id, user_id=user.id).first()


def _create_session_for_user(user: User) -> AssistantSession:
    session = AssistantSession(user_id=user.id)
    db.session.add(session)
    db.session.flush()
    return session


def _session_history(session: AssistantSession) -> List[Dict[str, str]]:
    messages: Sequence[AssistantMessage] = (
        AssistantMessage.query.filter_by(session_id=session.id)
        .order_by(AssistantMessage.created_at.asc(), AssistantMessage.id.asc())
        .all()
    )
    return [{"role": m.role, "content": m.content} for m in messages]


def _compose_document_context(session: AssistantSession, max_chars: int = 8000) -> Optional[str]:
    ready_documents = [
        doc
        for doc in session.documents
        if doc.status == "ready" and doc.extracted_text and doc.extracted_text.strip()
    ]
    if not ready_documents:
        return None
    sorted_docs = sorted(ready_documents, key=lambda d: d.created_at or date.min, reverse=True)
    remaining = max_chars
    sections: List[str] = []
    for doc in sorted_docs:
        snippet = doc.extracted_text.strip()
        if not snippet:
            continue
        if len(snippet) > remaining:
            snippet = snippet[: remaining]
        sections.append(
            f"Document: {doc.original_filename}\nContent excerpt:\n{snippet.strip()}"
        )
        remaining -= len(snippet)
        if remaining <= 0:
            break
    if not sections:
        return None
    return (
        "The user has attached the following documents. Use them as primary context when answering:\n\n"
        + "\n\n".join(sections)
    )


def _build_history(messages: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    valid = []
    for msg in messages or []:
        role = msg.get("role")
        content = msg.get("content")
        if role in {"user", "assistant", "system"} and isinstance(content, str):
            valid.append({"role": role, "content": content})
    return valid


def _normalize_message_text(message: str) -> str:
    if not message:
        return ""
    replacements = {
        "“": '"',
        "”": '"',
        "‘": "'",
        "’": "'",
    }
    normalized = message.translate(str.maketrans(replacements))
    lowered = normalized.lower()
    additions: Set[str] = set()
    for alias, replacement in MESSAGE_ALIAS_MAP.items():
        if alias in lowered and replacement not in lowered:
            additions.add(replacement)
    if additions:
        normalized = f"{normalized} {' '.join(sorted(additions))}"
    return normalized


def _safe_strip(value: Union[str, None], chars: Optional[str] = None) -> str:
    if not isinstance(value, str):
        return ""
    return value.strip(chars) if chars else value.strip()


def _contains_word(text: str, word: str) -> bool:
    if not text or not word:
        return False
    pattern = r"\b" + re.escape(word) + r"\b"
    return re.search(pattern, text, re.IGNORECASE) is not None


def _ensure_decimal(value: Any) -> Decimal:
    if isinstance(value, Decimal):
        return value
    try:
        return Decimal(str(value))
    except (InvalidOperation, TypeError, ValueError):
        return Decimal("0")


def _should_force_builtin(message: str) -> bool:
    lowered = (message or "").lower()
    if not lowered:
        return False
    if not any(keyword in lowered for keyword in SOFTWARE_KEYWORDS):
        return False
    return any(term in lowered for term in SENSITIVE_LICENSE_TERMS)


def _history_mentions_contract(history: List[Dict[str, str]], limit: int = 6) -> bool:
    if not history:
        return False
    for entry in reversed(history[-limit:]):
        content = _safe_strip(entry.get("content"))
        if not content:
            continue
        lowered = content.lower()
        if "contract" in lowered or "συμβ" in lowered or any(keyword in lowered for keyword in CONTRACT_KEYWORDS):
            return True
    return False


def _is_meaningful_builtin_reply(reply: Optional[str]) -> bool:
    if not reply:
        return False
    stripped = _safe_strip(reply)
    if not stripped:
        return False
    return stripped != BUILTIN_DEFAULT_RESPONSE


def _maybe_builtin_override(
    message: str,
    history: List[Dict[str, str]],
    user,
    reply: str,
    cached: Optional[str] = None,
) -> Optional[str]:
    if not _should_force_builtin(message):
        return None
    lowered_reply = (reply or "").lower()
    if not lowered_reply:
        return None
    if not any(marker in lowered_reply for marker in REFUSAL_MARKERS):
        return None
    builtin_reply = cached if _is_meaningful_builtin_reply(cached) else None
    if not builtin_reply:
        builtin_reply = _call_builtin(message, history, user)
    if _is_meaningful_builtin_reply(builtin_reply):
        current_app.logger.info("Assistant fallback: overriding LLM refusal with builtin software lookup.")
        return builtin_reply
    return None


def _should_replace_with_builtin(message: str, reply: str) -> bool:
    lowered_reply = (reply or "").lower()
    if not lowered_reply:
        return False
    lowered_message = (message or "").lower()
    domain_markers = {
        "ticket": ("no tickets", "no results for ticket"),
        "contract": ("no contracts", "no contract"),
        "hardware": ("no hardware", "no devices", "no asset"),
        "software": ("no software", "no license"),
        "network": ("no network", "no hosts"),
        "knowledge": ("no article", "no knowledge"),
        "contact": ("no contact", "no address book"),
    }
    for marker, phrases in domain_markers.items():
        if marker in lowered_message:
            for phrase in phrases:
                if phrase in lowered_reply:
                    return True
    return False


def _dispatch_module_query(tool_name: str, message: str, user) -> str:
    message = _normalize_message_text(message)
    lowered = message.lower()
    if tool_name == "query_tickets":
        response = _answer_ticket_query(message, lowered, user)
        return response or "No tickets matched those filters."
    if tool_name == "query_knowledge_base":
        response = _answer_knowledge_query(message)
        return response or "No knowledge items matched those filters."
    if tool_name == "query_software_inventory":
        response = _answer_software_query(message, lowered, user)
        return response or "No software assets matched those filters."
    if tool_name == "query_hardware_inventory":
        response = _answer_hardware_query(message, lowered, user)
        return response or "No hardware assets matched those filters."
    if tool_name == "query_network_inventory":
        response = _answer_network_query(message, lowered, user)
        return response or "No network records matched those filters."
    if tool_name == "query_contracts":
        response = _answer_contract_query(message, lowered, user)
        return response or "No contracts matched those filters."
    if tool_name == "query_address_book":
        response = _answer_address_book_query(message, lowered, user)
        return response or "No contacts matched those filters."
    return "Unsupported tool call."


def _session_response_payload(session: AssistantSession) -> Dict[str, Any]:
    refreshed = (
        AssistantSession.query.options(
            joinedload(AssistantSession.messages),
            joinedload(AssistantSession.documents),
        )
        .filter_by(id=session.id)
        .first()
    )
    target = refreshed or session
    return {
        "session": _serialize_session(target),
    }


@assistant_bp.route("/api/session", methods=["POST"])
@login_required
def api_create_or_fetch_session():
    if not is_feature_allowed("assistant_widget", current_user):
        abort(403)

    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id")
    session: Optional[AssistantSession] = None

    if session_id:
        session = _ensure_session_for_user(session_id, current_user)
        if not session:
            return (
                jsonify({"success": False, "message": _("Session not found.")}),
                404,
            )
        session.touch()
    else:
        session = _create_session_for_user(current_user)

    db.session.commit()
    return jsonify({"success": True, **_session_response_payload(session)})


@assistant_bp.route("/api/session/<int:session_id>", methods=["GET"])
@login_required
def api_get_session(session_id: int):
    if not is_feature_allowed("assistant_widget", current_user):
        abort(403)
    session = _ensure_session_for_user(session_id, current_user)
    if not session:
        return (
            jsonify({"success": False, "message": _("Session not found.")}),
            404,
        )
    session.touch()
    db.session.commit()
    return jsonify({"success": True, **_session_response_payload(session)})


@assistant_bp.route("/api/session/<int:session_id>/documents", methods=["POST"])
@login_required
def api_upload_document(session_id: int):
    if not is_feature_allowed("assistant_widget", current_user):
        abort(403)
    session = _ensure_session_for_user(session_id, current_user)
    if not session:
        return (
            jsonify({"success": False, "message": _("Session not found.")}),
            404,
        )

    if "file" not in request.files:
        return jsonify({"success": False, "message": _("No file provided.")}), 400
    file = request.files["file"]
    if not file or file.filename == "":
        return jsonify({"success": False, "message": _("Please choose a file.")}), 400

    filename = secure_filename(file.filename, allow_unicode=True)
    if not filename:
        return jsonify({"success": False, "message": _("Invalid filename.")}), 400

    upload_folder = _assistant_upload_folder()
    unique_name = f"{uuid.uuid4().hex}_{filename}"
    stored_path = os.path.join(upload_folder, unique_name)
    file.save(stored_path)

    mimetype = file.mimetype or mimetypes.guess_type(filename)[0] or ""
    try:
        file_size = os.path.getsize(stored_path)
    except OSError:
        file_size = None

    extracted_text = _extract_document_text(stored_path, mimetype)
    status = "ready" if extracted_text else "uploaded"

    document = AssistantDocument(
        session_id=session.id,
        user_id=current_user.id,
        original_filename=filename,
        stored_filename=unique_name,
        mimetype=mimetype,
        file_size=file_size,
        extracted_text=extracted_text,
        status=status,
    )
    db.session.add(document)
    session.touch()
    db.session.commit()
    return jsonify(
        {
            "success": True,
            "document": _serialize_document(document),
        }
    )


@assistant_bp.route(
    "/api/session/<int:session_id>/documents/<int:document_id>", methods=["DELETE"]
)
@login_required
def api_delete_document(session_id: int, document_id: int):
    if not is_feature_allowed("assistant_widget", current_user):
        abort(403)
    session = _ensure_session_for_user(session_id, current_user)
    if not session:
        return (
            jsonify({"success": False, "message": _("Session not found.")}),
            404,
        )
    document = AssistantDocument.query.filter_by(
        id=document_id, session_id=session.id
    ).first()
    if not document:
        return (
            jsonify({"success": False, "message": _("Document not found.")}),
            404,
        )
    upload_folder = _assistant_upload_folder()
    file_path = os.path.join(upload_folder, document.stored_filename)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except OSError:
            current_app.logger.warning("Failed to remove assistant document %s", file_path)
    db.session.delete(document)
    session.touch()
    db.session.commit()
    return jsonify({"success": True})


@assistant_bp.route(
    "/api/session/<int:session_id>/documents/<int:document_id>/download", methods=["GET"]
)
@login_required
def download_document(session_id: int, document_id: int):
    if not is_feature_allowed("assistant_widget", current_user):
        abort(403)
    session = _ensure_session_for_user(session_id, current_user)
    if not session:
        abort(404)
    document = AssistantDocument.query.filter_by(
        id=document_id, session_id=session.id
    ).first()
    if not document:
        abort(404)
    upload_folder = _assistant_upload_folder()
    return send_from_directory(
        upload_folder, document.stored_filename, as_attachment=True, download_name=document.original_filename
    )


@assistant_bp.route("/api/message", methods=["POST"])
@login_required
def api_message():
    if not is_feature_allowed("assistant_widget", current_user):
        abort(403)

    config = _load_config()
    if not config:
        return jsonify({"success": False, "message": _("Assistant is disabled.")}), 400

    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id")
    session = None
    if session_id:
        session = _ensure_session_for_user(session_id, current_user)
        if not session:
            return jsonify({"success": False, "message": _("Session not found.")}), 404
    else:
        session = _create_session_for_user(current_user)
        db.session.commit()

    message = _safe_strip(data.get("message"))
    if not message:
        return jsonify({"success": False, "message": _("Message cannot be empty.")}), 400

    normalized_message = _normalize_message_text(message)
    message = normalized_message

    system_prompt = config.system_prompt or DEFAULT_SYSTEM_PROMPT
    prior_history = _session_history(session)
    user_history_entry = {"role": "user", "content": message}
    history = prior_history + [user_history_entry]

    document_context = _compose_document_context(session)
    messages_for_model: List[Dict[str, str]] = []
    if system_prompt:
        messages_for_model.append({"role": "system", "content": system_prompt})
    messages_for_model.extend(prior_history)
    if document_context:
        messages_for_model.append({"role": "system", "content": document_context})
    messages_for_model.append(user_history_entry)

    reply: Optional[str] = None
    user_message_row = AssistantMessage(session_id=session.id, role="user", content=message)
    db.session.add(user_message_row)

    try:
        if config.provider == "chatgpt_hybrid":
            if not config.openai_api_key:
                return jsonify({"success": False, "message": _("OpenAI API key is not configured.")}), 400
            tool_context = {
                "user": current_user,
                "history": history,
                "latest_user_message": message,
            }
            reply = _call_openai(messages_for_model, config, allow_tools=True, tool_context=tool_context)
        elif config.provider == "openwebui":
            if not config.openwebui_base_url:
                return jsonify({"success": False, "message": _("OpenWebUI base URL is not configured.")}), 400
            tool_context = {
                "user": current_user,
                "history": history,
                "latest_user_message": message,
            }
            reply = _call_openwebui(messages_for_model, config, tool_context=tool_context)
        elif config.provider == "webhook":
            if not config.webhook_url:
                return jsonify({"success": False, "message": _("Webhook URL is not configured.")}), 400
            reply = _call_webhook(history, messages_for_model, config, system_prompt)
        else:
            return jsonify({"success": False, "message": _("Unsupported assistant provider configured.")}), 400
    except requests.RequestException as exc:
        return jsonify({"success": False, "message": _("Connection error: %(error)s", error=str(exc))}), 502
    except Exception as exc:  # pragma: no cover
        db.session.rollback()
        return jsonify({"success": False, "message": str(exc)}), 500

    if not reply:
        db.session.rollback()
        return jsonify({"success": False, "message": _("Assistant did not return a reply.")}), 502

    assistant_message_row = AssistantMessage(session_id=session.id, role="assistant", content=reply)
    db.session.add(assistant_message_row)
    session.touch()
    db.session.commit()

    history.append({"role": "assistant", "content": reply})
    updated_session = AssistantSession.query.get(session.id)
    session_payload = _session_response_payload(updated_session)
    return jsonify(
        {
            "success": True,
            "reply": reply,
            "history": history,
            **session_payload,
        }
    )


def _call_openai(
    messages: List[Dict[str, str]],
    config: AssistantConfig,
    allow_tools: bool,
    tool_context: Optional[Dict[str, Any]] = None,
    depth: int = 0,
) -> str:
    endpoint = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {config.openai_api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": config.openai_model or "gpt-3.5-turbo",
        "messages": messages,
    }
    if allow_tools:
        payload["tools"] = _build_openai_tool_spec()
        payload["tool_choice"] = "auto"
    response = requests.post(endpoint, headers=headers, json=payload, timeout=30)
    if response.status_code >= 400:
        raise RuntimeError(_("OpenAI error: %(status)s %(body)s", status=response.status_code, body=response.text))
    data = response.json()
    choices = data.get("choices") or []
    if not choices:
        return ""
    message = choices[0].get("message", {})
    tool_calls = message.get("tool_calls") if allow_tools else None
    if allow_tools and tool_calls:
        limit = _tool_call_limit()
        if limit >= 0 and depth >= limit:
            fallback = "Remote tool execution limit reached. Please respond directly with the information gathered so far."
            return _safe_strip(message.get("content", fallback)) or fallback
        new_messages = messages + [message]
        for call in tool_calls:
            result = _execute_tool_call(call, tool_context or {})
            new_messages.append(
                {
                    "role": "tool",
                    "tool_call_id": call.get("id") or "",
                    "content": result,
                }
            )
        return _call_openai(new_messages, config, allow_tools=True, tool_context=tool_context, depth=depth + 1)
    return _safe_strip(message.get("content", ""))


def _latest_user_message_from(messages: List[Dict[str, Any]]) -> str:
    for msg in reversed(messages or []):
        if msg.get("role") == "user":
            return _safe_strip(msg.get("content"))
    return ""


def _prepare_tool_context(
    tool_context: Optional[Dict[str, Any]],
    messages: List[Dict[str, str]],
) -> Dict[str, Any]:
    context = dict(tool_context or {})
    if not context.get("latest_user_message"):
        context["latest_user_message"] = _latest_user_message_from(messages)
    context.setdefault("history", messages)
    return context


def _tool_call_limit() -> int:
    """Return the maximum allowed tool recursion depth, or a negative value to disable."""
    try:
        limit = int(current_app.config.get("ASSISTANT_TOOL_CALL_DEPTH_LIMIT", 3))
    except (TypeError, ValueError):
        limit = 3
    return limit


def _build_openai_tool_spec() -> List[Dict[str, Any]]:
    specs: List[Dict[str, Any]] = []

    if current_app.config.get("MCP_ENABLED", True):
        for meta in _get_mcp_tool_metadata():
            parameters = meta.get("input_schema") or {"type": "object", "properties": {}}
            specs.append(
                {
                    "type": "function",
                    "function": {
                        "name": meta["name"],
                        "description": meta.get("description", "MCP tool"),
                        "parameters": parameters,
                    },
                }
            )

    return specs


def _execute_remote_openwebui_tool(
    config: AssistantConfig,
    tool_call: Dict[str, Any],
) -> str:
    base_url = (config.openwebui_base_url or "").rstrip("/")
    if not base_url:
        return "Remote tool execution failed: OpenWebUI base URL is missing."

    function = tool_call.get("function") or {}
    name = function.get("name")
    raw_args = function.get("arguments")
    try:
        args = json.loads(raw_args) if isinstance(raw_args, str) and raw_args else raw_args or {}
    except (ValueError, TypeError):
        args = {"raw": raw_args}

    endpoint = f"{base_url}/api/v1/tools/call"
    headers = {"Content-Type": "application/json"}
    if config.openwebui_api_key:
        headers["Authorization"] = f"Bearer {config.openwebui_api_key}"

    payload = {
        "name": name,
        "arguments": args,
        "tool_call_id": tool_call.get("id"),
    }

    try:
        response = requests.post(endpoint, headers=headers, json=payload, timeout=30)
    except requests.RequestException as exc:
        return f"Remote tool execution failed for {name}: {exc}"

    if response.status_code >= 400:
        return _safe_strip(
            _("Remote tool execution failed: %(status)s %(body)s", status=response.status_code, body=response.text)
        )

    try:
        data = response.json()
    except ValueError as exc:
        return f"Remote tool execution returned invalid JSON: {exc}"

    for key in ("result", "output", "reply", "content", "data"):
        if key in data:
            value = data[key]
            if isinstance(value, (dict, list)):
                return json.dumps(value)
            return _safe_strip(str(value))

    return "Remote tool execution succeeded but returned no result."


def _call_openwebui(
    messages: List[Dict[str, str]],
    config: AssistantConfig,
    tool_context: Optional[Dict[str, Any]] = None,
    depth: int = 0,
) -> str:
    base_url = (config.openwebui_base_url or "").rstrip("/")
    if not base_url:
        raise RuntimeError(_("OpenWebUI base URL is not configured."))

    endpoint = f"{base_url}/api/v1/chat/completions"
    headers = {"Content-Type": "application/json"}
    if config.openwebui_api_key:
        headers["Authorization"] = f"Bearer {config.openwebui_api_key}"

    payload = {
        "model": config.openwebui_model or "gpt-3.5-turbo",
        "messages": messages,
        "stream": False,
        "tools": _build_openai_tool_spec(),
        "tool_choice": "auto",
    }

    response = requests.post(endpoint, headers=headers, json=payload, timeout=30)
    if response.status_code >= 400:
        raise RuntimeError(_("OpenWebUI error: %(status)s %(body)s", status=response.status_code, body=response.text))
    data = response.json()
    choices = data.get("choices") or []
    if not choices:
        return ""
    message = choices[0].get("message", {})
    tool_calls = message.get("tool_calls") or []
    if tool_calls:
        limit = _tool_call_limit()
        if limit >= 0 and depth >= limit:
            fallback = "Remote tool execution limit reached. Please respond directly with the information gathered so far."
            return _safe_strip(message.get("content", fallback)) or fallback
        context = _prepare_tool_context(tool_context, messages)
        new_messages = messages + [message]
        local_tool_names = {tool["name"] for tool in _get_mcp_tool_metadata()}
        for call in tool_calls:
            name = (call.get("function") or {}).get("name") or ""
            if name in local_tool_names:
                result = _execute_tool_call(call, context)
            else:
                result = _execute_remote_openwebui_tool(config, call)
            new_messages.append(
                {
                    "role": "tool",
                    "tool_call_id": call.get("id") or name,
                    "content": result,
                }
            )
        return _call_openwebui(new_messages, config, tool_context=context, depth=depth + 1)

    content = _safe_strip(message.get("content", ""))
    if content:
        stripped = content.strip()
        if stripped.startswith("{") and stripped.endswith("}"):
            try:
                tool_candidate = json.loads(stripped)
            except ValueError:
                tool_candidate = None
            if isinstance(tool_candidate, dict) and "name" in tool_candidate:
                name = tool_candidate.get("name") or ""
                arguments = (
                    tool_candidate.get("args")
                    or tool_candidate.get("arguments")
                    or tool_candidate.get("parameters")
                    or {}
                )
                raw_arguments = json.dumps(arguments)
                tool_call = {
                    "id": tool_candidate.get("id") or name,
                    "function": {
                        "name": name,
                        "arguments": raw_arguments,
                    },
                }
                local_tool_names = {tool["name"] for tool in _get_mcp_tool_metadata()}
                context = _prepare_tool_context(tool_context, messages)
                if name in local_tool_names:
                    return _execute_tool_call(tool_call, context)
                return _execute_remote_openwebui_tool(config, tool_call)
        return content
    return ""


def _call_webhook(
    history: List[Dict[str, str]],
    messages_with_system: List[Dict[str, str]],
    config: AssistantConfig,
    system_prompt: str,
) -> str:
    payload = {
        "history": history,
        "messages_with_system": messages_with_system,
        "system_prompt": system_prompt,
        "current_user": {
            "id": current_user.id,
            "username": current_user.username,
            "role": current_user.role,
        },
        "latest": history[-1],
    }
    method = (config.webhook_method or "POST").upper()
    headers = {"Content-Type": "application/json"}
    headers.update(config.webhook_headers_data())
    if method == "GET":
        response = requests.get(config.webhook_url, params={"payload": json.dumps(payload)}, headers=headers, timeout=30)
    else:
        response = requests.request(method, config.webhook_url, headers=headers, json=payload, timeout=30)
    if response.status_code >= 400:
        raise RuntimeError(_("Webhook error: %(status)s %(body)s", status=response.status_code, body=response.text))
    try:
        data = response.json()
    except ValueError as exc:
        raise RuntimeError(_("Invalid webhook response: %(error)s", error=str(exc)))
    reply = data.get("reply")
    if isinstance(reply, str):
        return _safe_strip(reply)
    raise RuntimeError(_("Webhook response is missing a 'reply' field."))


def _format_mcp_result(tool: str, result: Any) -> str:
    """Produce a readable string from an MCP tool result."""

    if isinstance(result, dict):
        rows = result.get("rows")
        if isinstance(rows, list):
            if not rows:
                return _("Tool `%(tool)s` returned no rows.", tool=tool)
            lines = []
            knowledge_tool = tool in {"knowledge_recent_articles"}
            for idx, row in enumerate(rows, start=1):
                if isinstance(row, dict):
                    rendered = []
                    if knowledge_tool:
                        article_id = row.get("id") or row.get("article_id")
                        title = row.get("title")
                        if not title:
                            title = _("Untitled article") if not article_id else f"Article {article_id}"
                        link = row.get("url") or row.get("article_url")
                        if not link and article_id:
                            try:
                                link = url_for("knowledge.view_article", article_id=article_id)
                            except RuntimeError:
                                link = f"/knowledge/article/{article_id}"
                        rendered.append(title)
                        if link:
                            rendered.append(f"url: {link}")
                        summary = row.get("summary")
                        if summary:
                            rendered.append(f"summary: {summary}")
                        tags = row.get("tags")
                        if tags:
                            rendered.append(f"tags: {tags}")
                        attachments = row.get("attachments") or []
                        attachment_lines: List[str] = []
                        for attachment in attachments:
                            if not isinstance(attachment, dict):
                                continue
                            att_name = attachment.get("filename") or f"attachment {attachment.get('id')}"
                            att_url = attachment.get("download_url") or attachment.get("url")
                            if not att_url:
                                continue
                            size = attachment.get("file_size")
                            markdown = attachment.get("download_markdown")
                            if markdown:
                                line = f"- {markdown}"
                            else:
                                line = f"- [{att_name}]({att_url})"
                            if size:
                                line += f" ({size} bytes)"
                            attachment_lines.append(line)
                        if attachment_lines:
                            rendered.append("attachments:\n" + "\n".join(attachment_lines))
                    else:
                        rendered.append(", ".join(f"{key}: {row_val}" for key, row_val in row.items()))
                    formatted = " — ".join(part for part in rendered if part)
                else:
                    formatted = str(row)
                lines.append(f"{idx}. {formatted}")
            return _("Results from `%(tool)s`:\n%(rows)s", tool=tool, rows="\n".join(lines))
        return _("Tool `%(tool)s` response:\n%(payload)s", tool=tool, payload=json.dumps(result, indent=2, default=str))
    if isinstance(result, list):
        if not result:
            return _("Tool `%(tool)s` returned no data.", tool=tool)
        lines = [f"{idx}. {item}" for idx, item in enumerate(result, start=1)]
        return _("Results from `%(tool)s`:\n%(rows)s", tool=tool, rows="\n".join(lines))
    if result is None:
        return _("Tool `%(tool)s` completed with no data.", tool=tool)
    return _("Tool `%(tool)s` response: %(payload)s", tool=tool, payload=str(result))


def _invoke_mcp_tool(tool: str, arguments: Dict[str, Any]) -> Dict[str, Any]:
    base_url = current_app.config.get("MCP_BASE_URL")
    if not base_url:
        host = current_app.config.get("MCP_HOST", "127.0.0.1")
        port = current_app.config.get("MCP_PORT", 8081)
        base_url = f"http://{host}:{port}"
    endpoint = f"{base_url.rstrip('/')}/mcp/invoke"

    timeout = current_app.config.get("MCP_REQUEST_TIMEOUT_SECONDS", 10)
    response = requests.post(endpoint, json={"tool": tool, "arguments": arguments}, timeout=timeout)
    if response.status_code >= 400:
        try:
            detail = response.json()
        except ValueError:
            detail = response.text
        raise RuntimeError(
            _("MCP invocation failed (%(status)s): %(detail)s", status=response.status_code, detail=detail)
        )
    try:
        return response.json()
    except ValueError as exc:
        raise RuntimeError(_("Invalid response from MCP server: %(error)s", error=str(exc)))


def _get_mcp_tool_metadata() -> List[Dict[str, Any]]:
    cache = current_app.extensions.setdefault("mcp_server", {})
    cached = cache.get("tool_catalog")
    ttl = current_app.config.get("MCP_TOOL_CACHE_SECONDS", 300)
    now = time.time()
    if cached and now - cached.get("fetched_at", 0) < ttl:
        return cached.get("items", [])

    base_url = current_app.config.get("MCP_BASE_URL")
    if not base_url:
        host = current_app.config.get("MCP_HOST", "127.0.0.1")
        port = current_app.config.get("MCP_PORT", 8081)
        base_url = f"http://{host}:{port}"
    endpoint = f"{base_url.rstrip('/')}/mcp/tools"
    timeout = current_app.config.get("MCP_REQUEST_TIMEOUT_SECONDS", 10)
    try:
        response = requests.get(endpoint, timeout=timeout)
    except requests.RequestException:
        cache["tool_catalog"] = {"items": [], "fetched_at": now}
        return []
    if response.status_code >= 400:
        cache["tool_catalog"] = {"items": [], "fetched_at": now}
        return []
    try:
        data = response.json()
    except ValueError:
        cache["tool_catalog"] = {"items": [], "fetched_at": now}
        return []

    items = data.get("tools") or []
    cache["tool_catalog"] = {"items": items, "fetched_at": now}
    return items


def _execute_tool_call(tool_call: Dict[str, Any], context: Dict[str, Any]) -> str:
    function = tool_call.get("function") or {}
    name = function.get("name") or ""

    raw_args = function.get("arguments") or ""
    if isinstance(raw_args, dict):
        args = raw_args
    else:
        try:
            args = json.loads(raw_args) if raw_args else {}
        except (ValueError, TypeError):
            args = {"query": raw_args if isinstance(raw_args, str) else ""}

    mcp_tool_names = {tool["name"] for tool in _get_mcp_tool_metadata()}
    if name in mcp_tool_names:
        if not current_app.config.get("MCP_ENABLED", True):
            return "MCP server is disabled; cannot execute tool."
        try:
            response = _invoke_mcp_tool(name, args if isinstance(args, dict) else {})
        except Exception as exc:
            return f"MCP tool execution failed: {exc}"
        if isinstance(response, dict):
            return _format_mcp_result(name, response.get("data"))
        return str(response)

    return f"Unsupported tool call: {name}"


def _call_builtin(message: str, history: List[Dict[str, str]], user) -> str:
    """Answer simple operational questions by querying the local database."""

    lowered = message.lower()

    try:
        network_response = _answer_network_query(message, lowered, user)
    except Exception as exc:  # pragma: no cover
        current_app.logger.exception("Assistant network handler failed: %s", exc)
        network_response = ""
    if network_response:
        return network_response

    cross_response = _answer_cross_module_query(message, lowered, user)
    if cross_response:
        return cross_response

    has_backup = any(keyword in lowered for keyword in BACKUP_KEYWORDS)

    if has_backup:
        backup_response = _answer_backup_query(message, lowered)
        if backup_response:
            return backup_response

    has_ticket = any(keyword in lowered for keyword in TICKET_KEYWORDS)
    has_knowledge = any(keyword in lowered for keyword in KNOWLEDGE_KEYWORDS)
    has_hardware = any(keyword in lowered for keyword in HARDWARE_KEYWORDS)
    has_software = any(keyword in lowered for keyword in SOFTWARE_KEYWORDS)
    has_contracts = (
        "contract" in lowered
        or "συμβ" in lowered
        or any(keyword in lowered for keyword in CONTRACT_KEYWORDS)
    )
    if not has_contracts and history:
        if any(term in lowered for term in COST_KEYWORDS) and _history_mentions_contract(history):
            has_contracts = True
    has_contacts = (
        "contact" in lowered
        or "address book" in lowered
        or "επαφ" in lowered
        or any(keyword in lowered for keyword in ADDRESS_BOOK_KEYWORDS)
    )
    force_software = _should_force_builtin(message)

    if has_ticket:
        ticket_response = _answer_ticket_query(message, lowered, user)
        if ticket_response:
            return ticket_response

    if has_knowledge:
        knowledge_response = _answer_knowledge_query(message)
        if knowledge_response:
            return knowledge_response

    if has_software:
        software_response = _answer_software_query(message, lowered, user)
        if software_response:
            return software_response

    if has_hardware and not force_software:
        hardware_response = _answer_hardware_query(message, lowered, user)
        if hardware_response:
            return hardware_response

    if has_contracts:
        contract_response = _answer_contract_query(message, lowered, user)
        if contract_response:
            return contract_response

    if has_contacts:
        contacts_response = _answer_address_book_query(message, lowered, user)
        if contacts_response:
            return contacts_response

    return BUILTIN_DEFAULT_RESPONSE


def _answer_network_query(message: str, lowered: str, user) -> Optional[str]:
    has_network_intent = (
        any(_contains_word(lowered, term) for term in NETWORK_KEYWORDS)
        or CIDR_PATTERN.search(message) is not None
        or IP_ADDRESS_PATTERN.search(message) is not None
        or _contains_word(lowered, "host")
        or _contains_word(lowered, "hosts")
    )
    if not has_network_intent:
        return None

    candidate: Optional[Network] = None
    query = Network.query.options(joinedload(Network.hosts))
    networks_list = query.all()

    site_filter_value: Optional[str] = None
    site_match = SITE_FILTER_PATTERN.search(message)
    if site_match:
        site_filter_value = _safe_strip(site_match.group(1))
    else:
        at_site_match = re.search(r"networks?\s+at\s+(?:the\s+)?([a-z0-9_.\-\s]+)", lowered)
        if at_site_match:
            candidate_site = _safe_strip(at_site_match.group(1))
            if candidate_site:
                if candidate_site.endswith(" site"):
                    candidate_site = candidate_site[:-5].strip()
                site_filter_value = candidate_site

    if site_filter_value:
        normalized_site = site_filter_value.lower()
        site_matches = [
            net
            for net in networks_list
            if normalized_site in (net.site or "").lower()
        ]
        if site_matches:
            if len(site_matches) == 1:
                candidate = site_matches[0]
            else:
                display_label = site_filter_value
                lines = [f"Networks at site {display_label}:"]
                for net in sorted(site_matches, key=lambda n: (n.name or "", n.cidr or "")):
                    name = net.name or net.cidr or f"Network #{net.id}"
                    cidr = net.cidr or "n/a"
                    vlan = net.vlan or "n/a"
                    gateway = net.gateway or "n/a"
                    site_label = net.site or display_label
                    lines.append(
                        f"- {name} ({cidr}) | Site: {site_label} | VLAN: {vlan} | Gateway: {gateway}"
                    )
                return "\n".join(lines)
            # fall through with single match

    cidr_match = CIDR_PATTERN.search(message)

    if cidr_match:
        cidr_value = cidr_match.group(0)
        candidate = _find_network_by_cidr(networks_list, cidr_value)

    if not candidate:
        name_match = re.search(r"(?:network|subnet|vlan)\s+([a-z0-9_.\-\s]+)", lowered)
        if name_match:
            raw_name = name_match.group(1) or ""
            normalized = _normalize_network_name_hint(raw_name)
            if normalized:
                candidate = _match_network_by_normalized(networks_list, normalized)

    if not candidate:
        leading_match = re.search(r"([a-z0-9_.\-\s]+)\s+(?:network|subnet|vlan)", lowered)
        if leading_match:
            raw_name = leading_match.group(1) or ""
            normalized = _normalize_network_name_hint(raw_name)
            if normalized:
                candidate = _match_network_by_normalized(networks_list, normalized)

    if not candidate:
        for phrase in _extract_candidate_phrases(message):
            normalized = _normalize_network_name_hint(phrase)
            if not normalized:
                continue
            candidate = _match_network_by_normalized(networks_list, normalized)
            if candidate:
                break

    if not candidate and lowered.endswith(" network"):
        suffix = lowered[:-8].strip()
        normalized = _normalize_network_name_hint(suffix)
        if normalized:
            candidate = _match_network_by_normalized(networks_list, normalized)

    if not candidate:
        keywords = _extract_keywords(message, extra_stop=NETWORK_HINT_STOPWORDS)
        if keywords:
            for keyword in keywords[:3]:
                candidate = _match_network_by_normalized(networks_list, keyword)
                if candidate:
                    break

    if not candidate:
        user_tokens = _extract_user_tokens(message)
        host_tokens = _extract_network_host_tokens(message)
        user_usernames: List[str] = []
        if user and getattr(user, "username", None):
            user_usernames.append(user.username)
        if user_tokens or user_usernames or host_tokens:
            conditions = []
            search_tokens: List[str] = []
            for token in list(user_tokens) + list(host_tokens):
                cleaned = _safe_strip(token)
                if cleaned:
                    search_tokens.append(cleaned)

            for token in search_tokens:
                lowered_token = token.lower()
                like = f"%{lowered_token}%"
                conditions.append(
                    or_(
                        func.lower(NetworkHost.assigned_to).ilike(like),
                        func.lower(NetworkHost.hostname).ilike(like),
                        func.lower(NetworkHost.ip_address).ilike(like),
                        func.lower(NetworkHost.mac_address).ilike(like),
                        func.lower(NetworkHost.device_type).ilike(like),
                        func.lower(NetworkHost.description).ilike(like),
                    )
                )
            for uname in user_usernames:
                if uname:
                    conditions.append(func.lower(NetworkHost.assigned_to) == uname.lower())

            if conditions:
                host_query = NetworkHost.query.join(Network)
                host_results = (
                    host_query.filter(or_(*conditions))
                    .order_by(func.lower(Network.name).asc(), NetworkHost.ip_address.asc())
                    .limit(25)
                    .all()
                )
                if host_results:
                    lines = []
                    for host in host_results:
                        network_label = host.network.name or host.network.cidr or f"Network #{host.network_id}"
                        hostname = host.hostname or "—"
                        assigned = host.assigned_to or "Unassigned"
                        mac = host.mac_address or "—"
                        device_type = host.device_type or "—"
                        reserved_flag = "Reserved" if host.is_reserved else "Available"
                        description = host.description or "—"
                        lines.append(
                            f"{host.ip_address} — {network_label}; hostname {hostname}; MAC {mac}; type {device_type}; "
                            f"assigned to {assigned}; status {reserved_flag}; notes {description}"
                        )
                    if len(host_results) == 25:
                        lines.append("…limited to first 25 matches.")
                    return "Network matches:\n" + "\n".join(lines)
                return "No network hosts matched those details."

        if any(term in lowered for term in NETWORK_KEYWORDS) or "cidr" in lowered or "ip" in lowered:
            networks = sorted(
                networks_list,
                key=lambda net: (net.updated_at or date.min),
                reverse=True,
            )[:20]
            if not networks:
                return "No networks are currently registered in the database."
            lines = []
            for net in networks:
                host_count = len(net.hosts)
                site = net.site or "n/a"
                vlan = net.vlan or "n/a"
                gateway = net.gateway or "n/a"
                summary = (
                    f"{(net.name or 'Unnamed network')} ({net.cidr}) — network {net.network_address or 'n/a'}; "
                    f"broadcast {net.broadcast_address or 'n/a'}; hosts tracked {host_count}; "
                    f"gateway {gateway}; site {site}; VLAN {vlan}"
                )
                if net.description:
                    summary += f"; description {_safe_strip(net.description)}"
                if net.notes:
                    summary += f"; notes {_safe_strip(net.notes)}"
                lines.append(summary)
            if len(networks) == 20:
                lines.append("…showing latest 20 networks.")
            return "Tracked networks:\n" + "\n".join(lines)
        if cidr_match:
            return "I couldn't find that network in the database."
        return None

    ip_net = candidate.ip_network
    if not ip_net:
        return f"The network record '{candidate.name}' has an invalid CIDR ({candidate.cidr})."

    used_addresses = {host.ip_address for host in candidate.hosts if host.ip_address}
    reserved_addresses = {
        host.ip_address for host in candidate.hosts if host.ip_address and getattr(host, "is_reserved", False)
    }

    if any(token in lowered for token in ("list hosts", "show hosts", "all hosts", "hosts list", "detailed hosts", "host list", "ip status", "host status")) or _extract_network_host_tokens(message):
        return _format_network_hosts(candidate, message, lowered, user)

    requested_single = any(token in lowered for token in ("first", "single", "one"))
    limit = 1 if requested_single else 5
    available: List[str] = []
    for addr in ip_net.hosts():
        ip_str = str(addr)
        if ip_str in used_addresses:
            continue
        available.append(ip_str)
        if len(available) >= limit:
            break

    requested_pairs: List[Tuple[str, str]] = []
    host_count = len(candidate.hosts)

    def _add_pair(label: str, value: Optional[str]):
        requested_pairs.append((label, value or "n/a"))

    if "gateway" in lowered:
        _add_pair("Gateway", candidate.gateway)
    if "vlan" in lowered:
        _add_pair("VLAN", candidate.vlan)
    if any(term in lowered for term in ("site", "location")):
        _add_pair("Site", candidate.site)
    if "cidr" in lowered or "subnet" in lowered or "range" in lowered:
        _add_pair("CIDR", candidate.cidr)
    if "network address" in lowered or "network ip" in lowered or "base address" in lowered:
        _add_pair("Network", candidate.network_address)
    if "broadcast" in lowered:
        _add_pair("Broadcast", candidate.broadcast_address)
    if any(term in lowered for term in ("host count", "hosts count", "number of hosts", "host capacity")):
        _add_pair("Hosts tracked", str(host_count))
    if "description" in lowered:
        _add_pair("Description", _safe_strip(candidate.description))
    if "note" in lowered:
        _add_pair("Notes", _safe_strip(candidate.notes))
    if "reserved" in lowered:
        _add_pair("Reserved addresses tracked", str(len(reserved_addresses)))

    summary_lines = [f"Network {candidate.name} ({candidate.cidr})"]
    location_bits = []
    if candidate.site:
        location_bits.append(f"site {candidate.site}")
    if candidate.vlan:
        location_bits.append(f"VLAN {candidate.vlan}")
    if location_bits:
        summary_lines.append("; ".join(location_bits))
    if candidate.gateway:
        summary_lines.append(f"Gateway: {candidate.gateway}")
    summary_lines.append(f"Network: {candidate.network_address or 'n/a'}")
    summary_lines.append(f"Broadcast: {candidate.broadcast_address or 'n/a'}")
    if candidate.description:
        summary_lines.append(f"Description: {_safe_strip(candidate.description)}")
    if candidate.notes:
        summary_lines.append(f"Notes: {_safe_strip(candidate.notes)}")
    summary_lines.append(f"Assigned/reserved addresses tracked: {len(used_addresses)}")
    if reserved_addresses:
        preview = ", ".join(sorted(reserved_addresses)[:5])
        if len(reserved_addresses) > 5:
            preview += ", …"
        summary_lines.append(f"Reserved addresses: {preview}")

    if available:
        label = "Next available IP" if len(available) == 1 else "Next available IPs"
        summary_lines.append(f"{label}: {', '.join(available)}")
    else:
        summary_lines.append("No free addresses detected in this network.")

    if requested_pairs:
        detail_lines = [f"{candidate.name} ({candidate.cidr}) details:"]
        for label, value in requested_pairs:
            detail_lines.append(f"{label}: {value}")
        if "reserved" in lowered:
            detail_lines.append(f"Reserved addresses tracked: {len(reserved_addresses)}")
            if reserved_addresses:
                detail_lines.append("Reserved addresses: " + ", ".join(sorted(reserved_addresses)))
            else:
                detail_lines.append("No reserved addresses found in this network.")
        if "hosts" in lowered and not any(label.startswith("Hosts") for label, _ in requested_pairs):
            detail_lines.append(f"Hosts tracked: {host_count}")
        if available and any(term in lowered for term in ("available ip", "free ip", "next ip", "next available")):
            label = "Next available IP" if len(available) == 1 else "Next available IPs"
            detail_lines.append(f"{label}: {', '.join(available)}")
        return "\n".join(detail_lines)

    return "\n".join(summary_lines)


def _format_network_hosts(network: Network, message: str, lowered: str, user) -> str:
    hosts = list(network.hosts)
    if not hosts:
        return f"No hosts are registered under {network.name} ({network.cidr})."

    # Filter hosts by intent
    filters_applied: List[str] = []

    if "reserved" in lowered and "available" not in lowered:
        hosts = [h for h in hosts if h.is_reserved]
        filters_applied.append("reserved")
    elif "available" in lowered or "free" in lowered:
        hosts = [h for h in hosts if not h.is_reserved]
        filters_applied.append("available")

    user_tokens = _extract_user_tokens(message)
    if user_tokens:
        lowered_users = {token.lower() for token in user_tokens}
        hosts = [
            h for h in hosts if h.assigned_to and h.assigned_to.lower() in lowered_users
        ]
        filters_applied.append("assigned to specified user")
    elif any(term in lowered for term in ("assigned to me", "my hosts")) and user and getattr(user, "username", None):
        hosts = [h for h in hosts if h.assigned_to and h.assigned_to.lower() == user.username.lower()]
        filters_applied.append(f"assigned to {user.username}")

    host_tokens = _extract_network_host_tokens(message)
    if host_tokens:
        normalized_tokens = []
        for token in host_tokens:
            cleaned = _safe_strip(token, " \"'.,:;!")
            if cleaned:
                normalized_tokens.append(cleaned.lower())

        def matches_token(host: NetworkHost, token: str) -> bool:
            token_plain = token.replace(":", "").replace("-", "")
            mac_plain = (host.mac_address or "").replace(":", "").replace("-", "").lower()
            return any(
                (
                    host.ip_address and token in host.ip_address.lower(),
                    host.hostname and token in host.hostname.lower(),
                    mac_plain and token_plain in mac_plain,
                )
            )

        hosts = [h for h in hosts if any(matches_token(h, tok) for tok in normalized_tokens)]
        filters_applied.append("matching provided identifiers")

    device_match = re.search(r"(?:device|host)\\s+type\\s+([a-z0-9_.\\- ]+)", lowered)
    if device_match:
        device_value = _safe_strip(device_match.group(1))
        hosts = [
            h for h in hosts if h.device_type and device_value.lower() in h.device_type.lower()
        ]
        filters_applied.append(f"device type contains '{device_value}'")

    description_match = re.search(r"(?:notes?|description)\\s+(?:contains\\s+)?([a-z0-9_.\\- ]+)", lowered)
    if description_match:
        desc_value = _safe_strip(description_match.group(1))
        hosts = [
            h for h in hosts if h.description and desc_value.lower() in h.description.lower()
        ]
        filters_applied.append(f"description contains '{desc_value}'")

    if not hosts:
        filter_text = "; ".join(filters_applied) if filters_applied else "specified criteria"
        return f"No hosts matched {filter_text} under {network.name} ({network.cidr})."

    hosts = sorted(hosts, key=lambda h: (h.ip_address or "", h.hostname or ""))

    if len(hosts) == 1 and any(term in lowered for term in ("status", "reserved", "available")):
        host = hosts[0]
        status = "Reserved" if host.is_reserved else "Available"
        hostname = host.hostname or "—"
        assigned = host.assigned_to or "Unassigned"
        details = [f"IP {host.ip_address} is {status} in {network.name} ({network.cidr})."]
        details.append(f"Hostname: {hostname}.")
        details.append(f"Assigned to: {assigned}.")
        if host.device_type:
            details.append(f"Device type: {host.device_type}.")
        if host.description:
            details.append(f"Notes: {host.description}.")
        return " ".join(details)

    header = f"Hosts tracked for {network.name} ({network.cidr})"
    if filters_applied:
        header += f" — filters: {', '.join(filters_applied)}"
    lines = [header + ":"]
    for host in hosts[:25]:
        ip_addr = host.ip_address or "n/a"
        hostname = host.hostname or "—"
        mac = host.mac_address or "—"
        device_type = host.device_type or "—"
        assigned = host.assigned_to or "Unassigned"
        reserved = "Reserved" if host.is_reserved else "Available"
        description = host.description or "—"
        lines.append(
            f"{ip_addr} — hostname {hostname}; MAC {mac}; type {device_type}; assigned to {assigned}; "
            f"status {reserved}; description {description}"
        )
    if len(hosts) > 25:
        lines.append("…limited to first 25 hosts.")
    return "\n".join(lines)


def _answer_backup_query(message: str, lowered: str) -> Optional[str]:
    media_tokens = ("tape", "tapes", "cartridge", "cartridges", "lto", "disk", "disks", "drive", "drives", "storage media", "storage medium")
    has_media_context = any(token in lowered for token in media_tokens)
    if not has_media_context and "retention" not in lowered and "off-site" not in lowered and "offsite" not in lowered and "off site" not in lowered:
        if "backup" not in lowered:
            return None

    now = datetime.utcnow()

    medium_filter: Optional[str] = None
    if any(term in lowered for term in ("disk", "disks", "drive", "drives", "external disk", "external drive", "δίσκ", "δισκ")):
        medium_filter = "disk"
    elif any(term in lowered for term in ("tape", "tapes", "cartridge", "cartridges", "lto", "ταιν")):
        medium_filter = "tape"

    query = TapeCartridge.query.options(joinedload(TapeCartridge.current_location))
    if medium_filter:
        query = query.filter(TapeCartridge.medium_type == medium_filter)

    wants_overdue = any(term in lowered for term in ("expired", "overdue", "past due", "ληγ", "εκπρόθεσ", "εκπροθεσ"))
    wants_due = any(term in lowered for term in ("due", "within", "upcoming", "approaching", "λήγει", "ληγει", "σε", "πλησιάζει"))
    wants_retention = (
        "retention" in lowered
        or any(term in lowered for term in ("λήξη", "ληξη", "διατήρηση", "διατηρηση"))
        or wants_overdue
        or wants_due
    )

    def _location_label(medium: TapeCartridge) -> str:
        if medium.current_location:
            loc_type = medium.current_location.location_type or ""
            loc_label = BACKUP_LOCATION_LABELS.get(loc_type, loc_type.replace("_", " ").title() or "Unspecified")
            site_name = medium.current_location.site_name
            if site_name:
                return f"{loc_label} ({site_name})"
            return loc_label
        return "Unassigned"

    def _medium_label(medium: TapeCartridge) -> str:
        return "External disk" if medium.medium_type == "disk" else "Tape"

    if wants_retention:
        base = query.filter(TapeCartridge.retention_days.isnot(None), TapeCartridge.retention_until.isnot(None))
        if wants_overdue:
            results = (
                base.filter(TapeCartridge.retention_until <= now)
                .order_by(TapeCartridge.retention_until.asc())
                .all()
            )
            label = _("retention already expired")
        else:
            window = 7
            match = re.search(r"(\d+)\s*(?:day|days|ημέρ|ημερ)", lowered)
            if match:
                try:
                    window = max(int(match.group(1)), 1)
                except ValueError:
                    window = 7
            upper = now + timedelta(days=window)
            results = (
                base.filter(TapeCartridge.retention_until > now, TapeCartridge.retention_until <= upper)
                .order_by(TapeCartridge.retention_until.asc())
                .all()
            )
            label = _("retention due within %(days)s day(s)", days=window)

        if not results:
            return _("No storage media matched %(label)s.", label=label)

        lines = [
            _("Storage media with %(label)s: showing %(count)s item(s).", label=label, count=len(results))
        ]
        for medium in results[:25]:
            remaining = medium.retention_remaining_days()
            status_bits = []
            if remaining is not None:
                if remaining < 0:
                    status_bits.append(_("%(days)s day(s) overdue", days=abs(remaining)))
                else:
                    status_bits.append(_("%(days)s day(s) remaining", days=remaining))
            info = f"- {medium.barcode} ({_medium_label(medium)}) — "
            info += _("ends %(date)s", date=medium.retention_until.strftime("%Y-%m-%d %H:%M"))
            info += "; " + _location_label(medium)
            if status_bits:
                info += " (" + ", ".join(status_bits) + ")"
            lines.append(info)
        if len(results) > 25:
            lines.append("…" + _("limited to first 25 media."))
        return "\n".join(lines)

    wants_offsite = any(term in lowered for term in ("off-site", "off site", "offsite", "vault", "εκτός", "εκτος"))
    if wants_offsite:
        offsite_query = (
            query.join(TapeLocation, TapeLocation.id == TapeCartridge.current_location_id)
            .filter(TapeLocation.location_type == "off_site")
            .order_by(TapeCartridge.barcode.asc())
        )
        results = offsite_query.limit(50).all()
        if not results:
            return _("No storage media are currently marked as off-site.")
        lines = [_("Off-site storage media (%(count)s found):", count=len(results))]
        for medium in results:
            site_name = medium.current_location.site_name if medium.current_location else None
            site_text = f" — {site_name}" if site_name else ""
            lines.append(f"- {medium.barcode} ({_medium_label(medium)}){site_text}")
        return "\n".join(lines)

    # Default summary
    media = query.order_by(TapeCartridge.barcode.asc()).all()
    if not media:
        if medium_filter == "disk":
            return _("No external disks have been registered yet.")
        if medium_filter == "tape":
            return _("No tape cartridges have been registered yet.")
        return _("No storage media have been registered yet.")

    type_counts = Counter(m.medium_type or "tape" for m in media)
    location_counts = Counter(
        (m.current_location.location_type if m.current_location else "unassigned")
        for m in media
    )
    overdue = sum(1 for m in media if m.retention_days and m.retention_until and m.retention_until <= now)
    due_within_7 = sum(
        1
        for m in media
        if m.retention_days
        and m.retention_until
        and 0 < (m.retention_until - now).total_seconds() <= 7 * 24 * 3600
    )

    type_label = {
        "tape": _("tape cartridges"),
        "disk": _("external disks"),
    }
    lines = []
    if medium_filter:
        lines.append(_("Storage media summary (filtered by %(type)s):", type=type_label.get(medium_filter, medium_filter)))
    else:
        lines.append(_("Storage media summary:"))
    lines.append(_("Total media: %(count)s", count=len(media)))
    for key, count in type_counts.items():
        lines.append(_("- %(label)s: %(count)s", label=type_label.get(key, key.title()), count=count))
    for key, count in location_counts.items():
        loc_label = BACKUP_LOCATION_LABELS.get(key, "Unassigned" if key == "unassigned" else key.replace("_", " ").title())
        lines.append(_("- Location %(loc)s: %(count)s", loc=loc_label, count=count))
    lines.append(_("Retention overdue: %(count)s", count=overdue))
    lines.append(_("Retention due within 7 days: %(count)s", count=due_within_7))
    return "\n".join(lines)


def _answer_ticket_query(message: str, lowered: str, user) -> Optional[str]:
    id_match = TICKET_ID_PATTERN.search(message)
    if id_match:
        ticket_id_raw = id_match.group(1) or id_match.group(2)
        try:
            ticket_id = int(ticket_id_raw)
        except (TypeError, ValueError):
            ticket_id = None
        ticket = None
        if ticket_id is not None:
            ticket = (
                Ticket.query.options(joinedload(Ticket.assignee))
                .filter(Ticket.id == ticket_id)
                .first()
            )
        if not ticket:
            return f"Ticket #{ticket_id_raw} was not found."

        wants_comments = "comment" in lowered or "comments" in lowered
        wants_attachments = "attachment" in lowered or "attachments" in lowered
        wants_audit = "audit" in lowered or "history" in lowered or "log" in lowered

        if wants_comments:
            comments = (
                TicketComment.query.filter_by(ticket_id=ticket.id)
                .order_by(TicketComment.created_at.asc())
                .limit(50)
                .all()
            )
            if not comments:
                return f"Ticket #{ticket.id} has no recorded comments."
            lines = [
                f"{idx + 1}. {comment.user or 'Unknown'} — {comment.comment or '—'} "
                f"({comment.created_at.strftime('%Y-%m-%d %H:%M') if comment.created_at else 'n/a'})"
                for idx, comment in enumerate(comments)
            ]
            return "Comments for ticket #{0}:\n{1}".format(ticket.id, "\n".join(lines))

        if wants_attachments:
            attachments = (
                TicketAttachment.query.filter_by(ticket_id=ticket.id)
                .order_by(TicketAttachment.uploaded_at.asc())
                .limit(50)
                .all()
            )
            if not attachments:
                return f"Ticket #{ticket.id} has no attachments."
            lines = [
                f"{idx + 1}. {attachment.filename or attachment.filepath or 'Attachment'} "
                f"(uploaded by {attachment.uploaded_by or 'unknown'} on "
                f"{attachment.uploaded_at.strftime('%Y-%m-%d %H:%M') if attachment.uploaded_at else 'n/a'})"
                for idx, attachment in enumerate(attachments)
            ]
            return "Attachments for ticket #{0}:\n{1}".format(ticket.id, "\n".join(lines))

        if wants_audit:
            logs = (
                AuditLog.query.filter_by(ticket_id=ticket.id)
                .order_by(AuditLog.timestamp.desc())
                .limit(50)
                .all()
            )
            if not logs:
                return f"No audit log entries were found for ticket #{ticket.id}."
            lines = [
                f"{log.timestamp.strftime('%Y-%m-%d %H:%M') if log.timestamp else 'n/a'} — "
                f"{log.username or 'system'}: {log.action or '—'}"
                for log in logs
            ]
            return "Audit log for ticket #{0}:\n{1}".format(ticket.id, "\n".join(lines))

        assignee = ticket.assignee.username if ticket.assignee else "Unassigned"
        created = ticket.created_at.strftime("%Y-%m-%d %H:%M") if ticket.created_at else "—"
        updated = ticket.updated_at.strftime("%Y-%m-%d %H:%M") if ticket.updated_at else "—"
        return (
            f"Ticket #{ticket.id}: {ticket.subject}\n"
            f"Status: {ticket.status or 'Unknown'} | Priority: {ticket.priority or 'n/a'} | Department: {ticket.department or 'n/a'}\n"
            f"Assigned to: {assignee} | Created: {created} | Updated: {updated}"
        )

    base_query = Ticket.query.options(joinedload(Ticket.assignee))
    filters: List[str] = []
    results: List[Ticket] = []
    total: Optional[int] = None
    excluded_terms: Set[str] = set()
    explicit_status = False
    explicit_priority = False

    if any(token in lowered for token in ("assigned to me", "my ticket", "my tickets")) and user:
        base_query = base_query.filter(Ticket.assigned_to == user.id)
        filters.append(f"assigned to {user.username}")
        excluded_terms.add(user.username.lower())
    else:
        assignee_match = USER_REF_PATTERN.search(message)
        if assignee_match:
            target = _resolve_user_reference(assignee_match.group(1))
            if target:
                base_query = base_query.filter(Ticket.assigned_to == target.id)
                filters.append(f"assigned to {target.username}")
                excluded_terms.add(target.username.lower())

    if "unassigned" in lowered:
        base_query = base_query.filter(Ticket.assigned_to.is_(None))
        filters.append("unassigned")

    creator_match = CREATED_BY_PATTERN.search(message)
    if creator_match:
        creator = _resolve_user_reference(creator_match.group(1))
        if creator:
            base_query = base_query.filter(Ticket.created_by == creator.id)
            filters.append(f"created by {creator.username}")
            excluded_terms.add(creator.username.lower())

    if "created by me" in lowered and user:
        base_query = base_query.filter(Ticket.created_by == user.id)
        filters.append("created by you")
        excluded_terms.add(user.username.lower())

    dept_match = re.search(r"(?:department|dept|τμήμα|τμημα)\s+(?:is|=|:|στο|στην|στον|σε)?\s*([a-z0-9_.\-άέήίόύώ\s]+)", lowered)
    if dept_match:
        department_value = _safe_strip(dept_match.group(1))
        if department_value:
            department_value = re.split(r"\b(?:tickets?|incidents?|requests?)\b", department_value, 1)[0]
            department_value = _safe_strip(re.split(r"\b(?:with|and|or|,|where)\b", department_value, 1)[0])
            if department_value:
                base_query = base_query.filter(func.lower(Ticket.department) == department_value.lower())
                filters.append(f"department {department_value}")
                excluded_terms.add(department_value.lower())

    if "today" in lowered:
        today = date.today()
        if "updated" in lowered and "created" not in lowered:
            base_query = base_query.filter(func.date(Ticket.updated_at) == today)
            filters.append("updated today")
        else:
            base_query = base_query.filter(func.date(Ticket.created_at) == today)
            filters.append("created today")

    if "yesterday" in lowered:
        yesterday = date.today() - timedelta(days=1)
        base_query = base_query.filter(func.date(Ticket.created_at) == yesterday)
        filters.append("created yesterday")

    subject_terms = _extract_field_terms(message, SUBJECT_QUOTED_PATTERN, SUBJECT_SINGLE_QUOTED_PATTERN)
    for term in subject_terms:
        base_query = base_query.filter(Ticket.subject.ilike(f"%{term}%"))
        filters.append(f"subject contains \"{term}\"")
        excluded_terms.add(term.lower())

    description_terms = _extract_field_terms(
        message, DESCRIPTION_QUOTED_PATTERN, DESCRIPTION_SINGLE_QUOTED_PATTERN
    )
    for term in description_terms:
        base_query = base_query.filter(Ticket.description.ilike(f"%{term}%"))
        filters.append(f"description contains \"{term}\"")
        excluded_terms.add(term.lower())

    status_match = STATUS_PATTERN.search(message)
    if status_match:
        status_value = _safe_strip(status_match.group(1))
        if status_value:
            base_query = base_query.filter(func.lower(Ticket.status) == status_value.lower())
            filters.append(f"status {status_value}")
            excluded_terms.add(status_value.lower())
            explicit_status = True

    priority_match = PRIORITY_PATTERN.search(message)
    if priority_match:
        priority_value = _safe_strip(priority_match.group(1))
        if priority_value:
            base_query = base_query.filter(func.lower(Ticket.priority) == priority_value.lower())
            filters.append(f"priority {priority_value}")
            excluded_terms.add(priority_value.lower())
            explicit_priority = True

    if any(token in lowered for token in ("open", "pending", "progress", "unresolved")):
        if not explicit_status:
            base_query = base_query.filter(func.lower(Ticket.status).in_(OPEN_STATUS_VALUES))
            filters.append("status in open set")
    elif any(token in lowered for token in ("closed", "resolved", "done", "completed", "cancelled")):
        if not explicit_status:
            base_query = base_query.filter(func.lower(Ticket.status).in_(CLOSED_STATUS_VALUES))
            filters.append("status closed")

    if "overdue" in lowered:
        overdue_threshold = date.today() - timedelta(days=7)
        base_query = base_query.filter(Ticket.closed_at.is_(None))
        base_query = base_query.filter(func.date(Ticket.created_at) <= overdue_threshold)
        filters.append("overdue (>7 days open)")

    if "reopen" in lowered or "re-open" in lowered:
        recent_threshold = date.today() - timedelta(days=7)
        base_query = base_query.filter(Ticket.closed_at.isnot(None))
        base_query = base_query.filter(func.date(Ticket.closed_at) >= recent_threshold)
        filters.append("closed within last 7 days")

    if ("attachment" in lowered or "attachments" in lowered) and "with attachments" not in filters:
        base_query = base_query.filter(Ticket.attachments.any())
        filters.append("with attachments")

    if not explicit_priority:
        for level in PRIORITY_LEVELS:
            if f"{level} priority" in lowered or f"priority {level}" in lowered:
                base_query = base_query.filter(func.lower(Ticket.priority) == level)
                filters.append(f"priority {level}")
                excluded_terms.add(level)
                explicit_priority = True
                break

    created_between = CREATED_BETWEEN_PATTERN.search(message)
    if created_between:
        start = _parse_date_string(created_between.group(1))
        end = _parse_date_string(created_between.group(2))
        if start and end:
            base_query = base_query.filter(
                func.date(Ticket.created_at).between(start, end)
            )
            filters.append(f"created between {start.isoformat()} and {end.isoformat()}")
    else:
        generic_between = GENERIC_BETWEEN_PATTERN.search(lowered)
        if generic_between and "created" not in lowered and "closed" not in lowered and "updated" not in lowered:
            start = _parse_date_string(generic_between.group(1))
            end = _parse_date_string(generic_between.group(2))
            if start and end:
                base_query = base_query.filter(
                    func.date(Ticket.created_at).between(start, end)
                )
                filters.append(f"created between {start.isoformat()} and {end.isoformat()}")

    created_on = CREATED_ON_PATTERN.search(message)
    if created_on:
        target_date = _parse_date_string(created_on.group(1))
        if target_date:
            base_query = base_query.filter(func.date(Ticket.created_at) == target_date)
            filters.append(f"created on {target_date.isoformat()}")

    created_after = CREATED_AFTER_PATTERN.search(message)
    if created_after:
        target_date = _parse_date_string(created_after.group(1))
        if target_date:
            base_query = base_query.filter(func.date(Ticket.created_at) >= target_date)
            filters.append(f"created after {target_date.isoformat()}")

    created_before = CREATED_BEFORE_PATTERN.search(message)
    if created_before:
        target_date = _parse_date_string(created_before.group(1))
        if target_date:
            base_query = base_query.filter(func.date(Ticket.created_at) <= target_date)
            filters.append(f"created before {target_date.isoformat()}")

    if "closed today" in lowered:
        today = date.today()
        base_query = base_query.filter(Ticket.closed_at.isnot(None))
        base_query = base_query.filter(func.date(Ticket.closed_at) == today)
        filters.append("closed today")

    if "closed yesterday" in lowered:
        yesterday = date.today() - timedelta(days=1)
        base_query = base_query.filter(Ticket.closed_at.isnot(None))
        base_query = base_query.filter(func.date(Ticket.closed_at) == yesterday)
        filters.append("closed yesterday")

    closed_between = CLOSED_BETWEEN_PATTERN.search(message)
    if closed_between:
        start = _parse_date_string(closed_between.group(1))
        end = _parse_date_string(closed_between.group(2))
        if start and end:
            base_query = base_query.filter(Ticket.closed_at.isnot(None))
            base_query = base_query.filter(
                func.date(Ticket.closed_at).between(start, end)
            )
            filters.append(f"closed between {start.isoformat()} and {end.isoformat()}")

    closed_on = CLOSED_ON_PATTERN.search(message)
    if closed_on:
        target_date = _parse_date_string(closed_on.group(1))
        if target_date:
            base_query = base_query.filter(Ticket.closed_at.isnot(None))
            base_query = base_query.filter(func.date(Ticket.closed_at) == target_date)
            filters.append(f"closed on {target_date.isoformat()}")

    closed_after = CLOSED_AFTER_PATTERN.search(message)
    if closed_after:
        target_date = _parse_date_string(closed_after.group(1))
        if target_date:
            base_query = base_query.filter(Ticket.closed_at.isnot(None))
            base_query = base_query.filter(func.date(Ticket.closed_at) >= target_date)
            filters.append(f"closed after {target_date.isoformat()}")

    closed_before = CLOSED_BEFORE_PATTERN.search(message)
    if closed_before:
        target_date = _parse_date_string(closed_before.group(1))
        if target_date:
            base_query = base_query.filter(Ticket.closed_at.isnot(None))
            base_query = base_query.filter(func.date(Ticket.closed_at) <= target_date)
            filters.append(f"closed before {target_date.isoformat()}")

    need_total = "how many" in lowered or "count" in lowered
    used_content_filters = False
    phrases = []
    for phrase in _extract_candidate_phrases(message):
        phrase_lower = phrase.lower()
        if phrase_lower in excluded_terms:
            continue
        tokens = set(re.findall(r"[a-z0-9]+", phrase_lower))
        if tokens & excluded_terms:
            continue
        phrases.append(phrase)
    if phrases:
        used_content_filters = True
    for phrase in phrases:
        like = f"%{phrase}%"
        phrase_query = base_query.filter(
            or_(
                Ticket.subject.ilike(like),
                Ticket.description.ilike(like),
                Ticket.department.ilike(like),
            )
        )
        phrase_results = phrase_query.order_by(Ticket.created_at.desc()).limit(20).all()
        if phrase_results:
            results = phrase_results
            total = phrase_query.count() if need_total else len(results)
            break

    if not results:
        keywords = [
            keyword
            for keyword in _extract_keywords(message, extra_stop=TICKET_KEYWORDS | TICKET_TEXT_STOP)
            if keyword.lower() not in excluded_terms
        ]
        if keywords:
            used_content_filters = True
        keyword_query = base_query
        for keyword in keywords[:4]:
            like = f"%{keyword}%"
            keyword_query = keyword_query.filter(
                or_(
                    Ticket.subject.ilike(like),
                    Ticket.description.ilike(like),
                    Ticket.department.ilike(like),
                )
            )
        results = keyword_query.order_by(Ticket.created_at.desc()).limit(20).all()
        if results:
            total = keyword_query.count() if need_total else len(results)

    if not results:
        if filters or used_content_filters:
            return "No tickets matched those filters."
        fallback_query = base_query.order_by(Ticket.created_at.desc())
        results = fallback_query.limit(20).all()
        total = fallback_query.count() if need_total else len(results)

    if not results:
        if filters:
            return "No tickets matched those filters."
        return None

    lines = []
    for ticket in results:
        assignee = ticket.assignee.username if ticket.assignee else "Unassigned"
        created = ticket.created_at.strftime("%Y-%m-%d") if ticket.created_at else "—"
        line = (
            f"#{ticket.id} {ticket.subject} — {ticket.status or 'Unknown'}"
            f" | Priority {ticket.priority or 'n/a'} | Assigned to {assignee} | Created {created}"
        )
        lines.append(line)

    if total > len(results):
        lines.append(f"…and {total - len(results)} more matching tickets.")

    if filters:
        filter_text = ", ".join(filters)
        header = f"Found {total} ticket(s) ({filter_text})."
    else:
        header = f"Found {total} ticket(s)."

    return "\n".join([header] + lines)


def _answer_knowledge_query(message: str) -> Optional[str]:
    lowered = message.lower()
    filters: List[str] = []

    article_id = None
    article_match = re.search(r"article\s+#?(\d+)", lowered)
    if article_match:
        try:
            article_id = int(article_match.group(1))
        except (TypeError, ValueError):
            article_id = None

    if article_id is not None:
        if "latest version" in lowered:
            version = (
                KnowledgeArticleVersion.query.filter_by(article_id=article_id)
                .order_by(KnowledgeArticleVersion.version_number.desc())
                .first()
            )
            if not version:
                return f"No version history recorded for article {article_id}."
            created_at = version.created_at.strftime("%Y-%m-%d %H:%M") if version.created_at else "n/a"
            return (
                f"Article {article_id} latest version v{version.version_number}: {version.title}\n"
                f"Created by user #{version.created_by} on {created_at}."
            )

        if "version history" in lowered:
            versions = (
                KnowledgeArticleVersion.query.filter_by(article_id=article_id)
                .order_by(KnowledgeArticleVersion.version_number.desc())
                .limit(15)
                .all()
            )
            if not versions:
                return f"No version history recorded for article {article_id}."
            lines = [
                f"v{version.version_number} — {version.title} (created {version.created_at.strftime('%Y-%m-%d %H:%M') if version.created_at else 'n/a'} by user #{version.created_by})"
                for version in versions
            ]
            return f"Version history for article {article_id}:\n" + "\n".join(lines)

        if "attachment" in lowered or "attachments" in lowered:
            attachments = (
                KnowledgeAttachment.query.filter_by(article_id=article_id)
                .order_by(KnowledgeAttachment.uploaded_at.desc())
                .limit(25)
                .all()
            )
            if not attachments:
                return f"No attachments are linked to article {article_id}."
            lines = [
                f"{att.original_filename} ({att.mimetype or 'unknown'}; uploaded {att.uploaded_at.strftime('%Y-%m-%d %H:%M') if att.uploaded_at else 'n/a'})"
                for att in attachments
            ]
            return f"Attachments for article {article_id}:\n" + "\n".join(lines)

    if "draft vs" in lowered or "draft versus" in lowered:
        published = KnowledgeArticle.query.filter(KnowledgeArticle.is_published.is_(True)).count()
        drafts = KnowledgeArticle.query.filter(KnowledgeArticle.is_published.is_(False)).count()
        return f"Knowledge base totals — Published: {published}, Draft: {drafts}."

    query = (
        KnowledgeArticle.query.outerjoin(
            KnowledgeAttachment,
            KnowledgeAttachment.article_id == KnowledgeArticle.id,
        )
        .options(joinedload(KnowledgeArticle.attachments))
        .distinct()
    )

    if "published" in lowered and "list" in lowered:
        query = query.filter(KnowledgeArticle.is_published.is_(True))
        filters.append("published")
    elif "draft" in lowered and ("list" in lowered or "show" in lowered):
        query = query.filter(KnowledgeArticle.is_published.is_(False))
        filters.append("draft")
    else:
        query = query.filter(KnowledgeArticle.is_published.is_(True))

    tag_match = TAG_PATTERN.search(message)
    if tag_match:
        tag_value = _safe_strip(tag_match.group(1))
        if tag_value:
            like = f"%{tag_value}%"
            query = query.filter(KnowledgeArticle.tags.ilike(like))
            filters.append(f"tag {tag_value}")

    dept_match = DEPARTMENT_PATTERN.search(message)
    if dept_match:
        dept_value = _safe_strip(dept_match.group(1))
        if dept_value:
            like = f"%{dept_value}%"
            query = query.filter(
                or_(
                    KnowledgeArticle.category.ilike(like),
                    KnowledgeArticle.tags.ilike(like),
                    KnowledgeArticle.summary.ilike(like),
                )
            )
            filters.append(f"context {dept_value}")

    if "recently updated" in lowered or ("recent" in lowered and "updated" in lowered):
        recent_threshold = datetime.utcnow() - timedelta(days=14)
        query = query.filter(KnowledgeArticle.updated_at >= recent_threshold)
        filters.append("updated last 14 days")

    if "updated" in lowered:
        between_match = GENERIC_BETWEEN_PATTERN.search(lowered)
        if between_match:
            start = _parse_date_string(between_match.group(1))
            end = _parse_date_string(between_match.group(2))
            if start and end:
                query = query.filter(
                    KnowledgeArticle.updated_at.isnot(None),
                    func.date(KnowledgeArticle.updated_at).between(start, end),
                )
                filters.append(f"updated between {start} and {end}")

    keywords = _extract_keywords(message, extra_stop={"knowledge", "article", "articles", "guide", "manual", "procedure"})
    keyword_conditions = []
    for keyword in keywords[:6]:
        like = f"%{keyword}%"
        keyword_conditions.append(
            or_(
                KnowledgeArticle.title.ilike(like),
                KnowledgeArticle.summary.ilike(like),
                KnowledgeArticle.tags.ilike(like),
                KnowledgeArticle.content.ilike(like),
                KnowledgeAttachment.original_filename.ilike(like),
                KnowledgeAttachment.extracted_text.ilike(like),
            )
        )

    if keyword_conditions:
        query = query.filter(and_(*keyword_conditions))

    if not filters and not keyword_conditions and article_id is None:
        # no clear intent
        keywords_raw = _extract_keywords(message)
        if not keywords_raw:
            return None

    results = query.order_by(KnowledgeArticle.updated_at.desc()).limit(10).all()
    if not results:
        if filters or keyword_conditions:
            return "No knowledge items matched those filters."
        return None

    lines = []
    for article in results:
        updated = article.updated_at.strftime("%Y-%m-%d") if article.updated_at else "—"
        tags = article.tags or "n/a"
        attachment_count = len(article.attachments)
        attachment_part = f" | attachments: {attachment_count}" if attachment_count else ""
        try:
            link = url_for("knowledge.view_article", article_id=article.id, _external=True)
        except RuntimeError:
            link = ""
        link_part = f" | link: <{link}>" if link else ""
        status_part = "draft" if not article.is_published else "published"
        lines.append(
            f"#{article.id} {article.title} — {status_part}; tags: {tags}; updated {updated}{attachment_part}{link_part}"
        )

    header = "Knowledge base matches"
    if filters:
        header += f" ({', '.join(filters)})"
    return header + ":\n" + "\n".join(lines)


def _answer_contract_query(message: str, lowered: str, user) -> Optional[str]:
    base_query = Contract.query.options(joinedload(Contract.owner))
    filters: List[str] = []
    need_total = "how many" in lowered or "count" in lowered
    support_terms = ("support", "phone", "email", "contact", "τηλεφων", "επικοινων", "email")
    show_support = any(term in lowered for term in support_terms)
    suppressed_keywords: Set[str] = set()

    if "active" in lowered or "ενεργ" in lowered:
        base_query = base_query.filter(func.lower(Contract.status) == "active")
        filters.append("status active")
        suppressed_keywords.add("active")

    if "auto-renew" in lowered or "auto renew" in lowered:
        base_query = base_query.filter(Contract.auto_renew.is_(True))
        filters.append("auto-renew")
        suppressed_keywords.update({"autorenew", "auto", "renew"})

    if show_support:
        suppressed_keywords.add("support")

    vendor_value: Optional[str] = None
    normalized_vendor: Optional[str] = None
    vendor_match = VENDOR_PATTERN.search(message)
    if vendor_match:
        vendor_value = _safe_strip(vendor_match.group(1))
    else:
        inferred_match = CONTRACT_FROM_PATTERN.search(lowered)
        if inferred_match:
            vendor_value = _safe_strip(inferred_match.group(1))
        if not vendor_value:
            for_match = CONTRACT_FOR_PATTERN.search(lowered)
            if for_match:
                vendor_value = _safe_strip(for_match.group(1))
        if not vendor_value:
            trailing_match = VENDOR_TRAILING_PATTERN.search(message)
            if trailing_match:
                vendor_value = _safe_strip(trailing_match.group(1))
        if not vendor_value:
            candidate_phrases = _extract_candidate_phrases(message)
            for phrase in candidate_phrases:
                raw_tokens = [tok for tok in re.split(r"[\s,]+", phrase) if tok]
                filtered_tokens: List[str] = []
                for raw in raw_tokens:
                    cleaned_token = _safe_strip(re.sub(r"[^a-zA-Z0-9_.\-άέήίόύώ]", "", raw))
                    if not cleaned_token:
                        continue
                    if cleaned_token.lower() in CONTRACT_VENDOR_STOPWORDS:
                        continue
                    filtered_tokens.append(cleaned_token)
                if filtered_tokens:
                    vendor_value = " ".join(filtered_tokens)
                    break
        if not vendor_value:
            direct_match = re.search(
                r"contracts?\s+(?!for\b|and\b|with\b|about\b|regarding\b)([a-z0-9_.\-άέήίόύώ]+)",
                message,
                re.IGNORECASE,
            )
            if direct_match:
                candidate = _safe_strip(direct_match.group(1))
                if candidate and candidate.lower() not in CONTRACT_VENDOR_STOPWORDS:
                    vendor_value = candidate
                    filters.append(f"vendor token {vendor_value}")
            if not vendor_value:
                sequential_tokens = re.findall(r"[A-Za-z0-9_.\-άέήίόύώ]+", message)
                token_count = len(sequential_tokens)
                idx = 0
                while idx < token_count:
                    token = sequential_tokens[idx]
                    if token.lower() in {"contract", "contracts"}:
                        j = idx + 1
                        vendor_tokens: List[str] = []
                        while j < token_count:
                            next_token = _safe_strip(sequential_tokens[j])
                            j += 1
                            if not next_token:
                                continue
                            lowered_next = next_token.lower()
                            if lowered_next in CONTRACT_VENDOR_STOPWORDS:
                                if vendor_tokens and lowered_next in {"contracts", "contract"}:
                                    break
                                if lowered_next in {"and", "with"}:
                                    continue
                                break
                            vendor_tokens.append(next_token)
                        if vendor_tokens:
                            vendor_value = " ".join(vendor_tokens)
                            filters.append(f"vendor tokens {vendor_value}")
                            break
                    idx += 1
    if vendor_value:
        vendor_value = re.split(r"\b(?:contracts?|support|and|,|with|phone|email|contact|vendor)\b", vendor_value, 1)[0]
        vendor_value = _safe_strip(vendor_value)
        if vendor_value:
            like_pattern = f"%{vendor_value}%"
            base_query = base_query.filter(
                or_(
                    func.trim(Contract.vendor).ilike(like_pattern),
                    Contract.name.ilike(like_pattern),
                )
            )
            filters.append(f"vendor ~ {vendor_value}")
            suppressed_keywords.add(vendor_value.lower())
            normalized_vendor = vendor_value.lower()
    elif filters and any(entry.startswith("vendor tokens") or entry.startswith("vendor token") for entry in filters):
        dynamic_terms = [entry.split(" ", 2)[-1] for entry in filters if entry.startswith("vendor")]
        if dynamic_terms:
            like_filters = [
                or_(
                    func.trim(Contract.vendor).ilike(f"%{term}%"),
                    Contract.name.ilike(f"%{term}%"),
                )
                for term in dynamic_terms
            ]
            base_query = base_query.filter(or_(*like_filters))

    number_match = CONTRACT_NUMBER_PATTERN.search(message)
    if number_match:
        number_value = _safe_strip(number_match.group(1))
        if number_value:
            invalid_tokens = {"from", "for", "by", "with"}
            lower_number = number_value.lower()
            if lower_number in invalid_tokens or (
                normalized_vendor
                and (
                    lower_number == normalized_vendor
                    or normalized_vendor.startswith(lower_number)
                    or lower_number.startswith(normalized_vendor)
                )
            ):
                number_value = ""
            elif len(number_value) <= 2 and not re.search(r"[0-9]", number_value):
                number_value = ""
        if number_value:
            contract = (
                base_query.filter(func.lower(Contract.contract_number) == number_value.lower())
                .order_by(Contract.id.asc())
                .first()
            )
            if not contract:
                return f"Contract number {number_value} was not found."
            return _format_contract_detail(contract, include_support=True)

    type_match = CONTRACT_TYPE_PATTERN.search(message)
    if type_match:
        type_value = _safe_strip(type_match.group(1))
        if type_value:
            base_query = base_query.filter(func.lower(func.trim(Contract.contract_type)) == type_value.lower())
            filters.append(f"type {type_value}")
            suppressed_keywords.add(type_value.lower())

    owner_filtered = False
    if "owner" in lowered or "υπεύθ" in lowered:
        owner_match = USER_REF_PATTERN.search(message)
        if owner_match:
            owner_user = _resolve_user_reference(owner_match.group(1))
            if owner_user:
                base_query = base_query.filter(Contract.owner_id == owner_user.id)
                filters.append(f"owner {owner_user.username}")
                suppressed_keywords.add(owner_user.username.lower())
                owner_filtered = True
    if not owner_filtered and user and any(
        phrase in lowered for phrase in ("my contract", "my contracts", "assigned to me", "for me")
    ):
        base_query = base_query.filter(Contract.owner_id == user.id)
        filters.append(f"owner {user.username}")
        suppressed_keywords.add(user.username.lower())

    if "renewal" in lowered or "renewals" in lowered:
        date_by_match = DATE_BY_PATTERN.search(lowered)
        if date_by_match:
            target_date = _parse_date_string(date_by_match.group(1))
            if target_date:
                base_query = base_query.filter(Contract.renewal_date.isnot(None))
                base_query = base_query.filter(Contract.renewal_date <= target_date)
                filters.append(f"renewal by {target_date.isoformat()}")
                suppressed_keywords.add(str(target_date.year))

    if any(term in lowered for term in ("ending", "end date", "λήγουν", "ληγουν")):
        between_match = GENERIC_BETWEEN_PATTERN.search(lowered)
        if between_match:
            start = _parse_date_string(between_match.group(1))
            end = _parse_date_string(between_match.group(2))
            if start and end:
                base_query = base_query.filter(Contract.end_date.isnot(None))
                base_query = base_query.filter(Contract.end_date.between(start, end))
                filters.append(f"ending between {start.isoformat()} and {end.isoformat()}")
                suppressed_keywords.update({str(start.year), str(end.year)})
        else:
            end_by_match = DATE_BY_PATTERN.search(lowered)
            if end_by_match:
                target_date = _parse_date_string(end_by_match.group(1))
                if target_date:
                    base_query = base_query.filter(Contract.end_date.isnot(None))
                    base_query = base_query.filter(Contract.end_date <= target_date)
                    filters.append(f"ending by {target_date.isoformat()}")
                    suppressed_keywords.add(str(target_date.year))

    if any(term in lowered for term in ("high-value", "high value", "over", "άνω των", "ανω των")):
        amount_match = AMOUNT_THRESHOLD_PATTERN.search(lowered)
        if amount_match:
            amount_value = _parse_decimal(amount_match.group(1))
            if amount_value is not None:
                base_query = base_query.filter(Contract.value.isnot(None))
                base_query = base_query.filter(Contract.value >= amount_value)
                filters.append(f"value ≥ {amount_value}")
                digits_only = re.sub(r"\D", "", str(amount_value))
                if digits_only:
                    suppressed_keywords.add(digits_only)

    wants_summary = any(
        phrase in lowered
        for phrase in ("summarize", "summary", "aggregate", "rollup", "total value", "σύνολο", "συνολικά")
    )
    if wants_summary:
        summary_query = base_query.enable_eagerloads(False)
        vendor_expr = func.coalesce(func.nullif(func.trim(Contract.vendor), ""), literal(_("Unspecified vendor")))
        currency_expr = func.coalesce(func.nullif(func.trim(Contract.currency), ""), literal(_("Unspecified currency")))
        total_expr = func.coalesce(func.sum(func.coalesce(Contract.value, 0)), literal(0))
        count_expr = func.count(Contract.id)
        summary_rows = (
            summary_query.with_entities(
                vendor_expr.label("vendor"),
                currency_expr.label("currency"),
                total_expr.label("total_value"),
                count_expr.label("contract_count"),
            )
            .group_by(vendor_expr, currency_expr)
            .order_by(vendor_expr.asc(), currency_expr.asc())
            .all()
        )
        if not summary_rows:
            header = _("Overall contract value")
            overall_line = f"{header}: 0.00 across 0 contract(s)"
            if filters:
                overall_line += f" ({', '.join(filters)})"
            return overall_line
        lines: List[str] = []
        overall_value = Decimal("0")
        overall_count = 0
        for vendor_label, currency_label, total_value, contract_count in summary_rows:
            vendor_display = vendor_label or _("Unspecified vendor")
            currency_display = currency_label or _("Unspecified currency")
            amount = _ensure_decimal(total_value)
            overall_value += amount
            overall_count += contract_count or 0
            try:
                formatted_amount = f"{amount:,.2f}"
            except (InvalidOperation, TypeError, ValueError):
                formatted_amount = str(total_value)
            lines.append(
                f"{vendor_display} — {currency_display}: total {formatted_amount} across {contract_count} contract(s)"
            )
        try:
            formatted_overall = f"{overall_value:,.2f}"
        except (InvalidOperation, TypeError, ValueError):
            formatted_overall = str(overall_value)
        summary_header = _("Overall contract value")
        overall_line = f"{summary_header}: {formatted_overall} across {overall_count} contract(s)"

        if "vendor" in lowered or "currency" in lowered:
            header = _("Total contract value by vendor and currency")
        else:
            header = _("Total contract value breakdown by vendor and currency")
        if filters:
            header += f" ({', '.join(filters)})"
        return overall_line + "\n" + header + ":\n" + "\n".join(lines)

    keywords = _extract_keywords(
        message,
        extra_stop={"contract", "contracts", "renewal", "renewals", "search", "vendor", "vendors"},
    )
    keywords = [
        keyword
        for keyword in keywords
        if keyword not in suppressed_keywords
    ]
    keyword_conditions = []
    for keyword in keywords[:5]:
        like = f"%{keyword}%"
        keyword_conditions.append(
            or_(
                Contract.name.ilike(like),
                Contract.vendor.ilike(like),
                Contract.coverage_scope.ilike(like),
                Contract.notes.ilike(like),
            )
        )

    filtered_query = base_query.filter(and_(*keyword_conditions)) if keyword_conditions else base_query
    results = filtered_query.order_by(Contract.end_date.asc(), Contract.name.asc()).limit(20).all()
    total = filtered_query.count() if need_total else len(results)

    if not results:
        if filters or keyword_conditions:
            return "No contracts matched those filters."
        return None

    lines = []
    for contract in results:
        end_date = contract.end_date.strftime("%Y-%m-%d") if contract.end_date else "n/a"
        renewal = contract.renewal_date.strftime("%Y-%m-%d") if contract.renewal_date else "n/a"
        owner_name = contract.owner.username if contract.owner else "n/a"
        support_info = _support_contact_snippet(contract) if show_support else ""
        value_display = "n/a"
        if contract.value is not None:
            amount = _ensure_decimal(contract.value)
            try:
                formatted_amount = f"{amount:,.2f}"
            except (InvalidOperation, TypeError, ValueError):
                formatted_amount = str(contract.value)
            currency_part = f" {contract.currency}" if contract.currency else ""
            value_display = f"{formatted_amount}{currency_part}".strip()
        line = (
            f"{contract.name} — {contract.contract_type}; vendor {contract.vendor or 'n/a'}; "
            f"status {contract.status or 'n/a'}; end {end_date}; renewal {renewal}; "
            f"auto-renew {'yes' if contract.auto_renew else 'no'}; owner {owner_name}; value {value_display}"
        )
        if support_info:
            line += f"; support {support_info}"
        lines.append(line)

    header = f"Found {total} contract(s)"
    if filters:
        header += f" ({', '.join(filters)})"
    return header + ":\n" + "\n".join(lines)


def _answer_address_book_query(message: str, lowered: str, user) -> Optional[str]:
    base_query = AddressBookEntry.query
    filters: List[str] = []
    need_total = "how many" in lowered or "count" in lowered

    if "vendor" in lowered and "contact" in lowered:
        base_query = base_query.filter(func.lower(AddressBookEntry.category) == "vendor")
        filters.append("category Vendor")
    elif "partner" in lowered:
        base_query = base_query.filter(func.lower(AddressBookEntry.category) == "partner")
        filters.append("category Partner")
    elif "customer" in lowered:
        base_query = base_query.filter(func.lower(AddressBookEntry.category) == "customer")
        filters.append("category Customer")

    company_match = COMPANY_PATTERN.search(message)
    if company_match:
        company_value = _safe_strip(company_match.group(1))
        if company_value:
            base_query = base_query.filter(func.lower(AddressBookEntry.company) == company_value.lower())
            filters.append(f"company {company_value}")

    dept_match = DEPARTMENT_PATTERN.search(message)
    if dept_match:
        dept_value = _safe_strip(dept_match.group(1))
        if dept_value:
            base_query = base_query.filter(func.lower(AddressBookEntry.department) == dept_value.lower())
            filters.append(f"department {dept_value}")

    city_match = CITY_PATTERN.search(message)
    if city_match:
        city_value = _safe_strip(city_match.group(1))
        if city_value:
            base_query = base_query.filter(func.lower(AddressBookEntry.city) == city_value.lower())
            filters.append(f"city {city_value}")

    tag_match = TAG_PATTERN.search(message)
    if tag_match:
        tag_value = _safe_strip(tag_match.group(1))
        if tag_value:
            base_query = base_query.filter(AddressBookEntry.tags.ilike(f"%{tag_value}%"))
            filters.append(f"tag {tag_value}")

    domain_match = EMAIL_DOMAIN_PATTERN.search(message)
    if domain_match:
        domain_value = domain_match.group(1).lstrip("@")
        base_query = base_query.filter(
            or_(
                AddressBookEntry.email.ilike(f"%@{domain_value}%"),
                AddressBookEntry.website.ilike(f"%{domain_value}%"),
            )
        )
        filters.append(f"domain {domain_value}")

    phone_match = PHONE_PATTERN.search(message)
    if phone_match:
        phone_value = re.sub(r"[^0-9+]", "", phone_match.group(1) or "")
        if phone_value:
            base_query = base_query.filter(
                or_(
                    AddressBookEntry.phone.ilike(f"%{phone_value}%"),
                    AddressBookEntry.mobile.ilike(f"%{phone_value}%"),
                )
            )
            filters.append(f"phone {phone_value}")

    detail_candidates: List[str] = []
    contact_match = CONTACT_NAME_PATTERN.search(message)
    if contact_match:
        detail_candidates.append(contact_match.group(1))
    if "contact details" in lowered or "show contact" in lowered or "find contact" in lowered:
        detail_candidates.extend(_extract_candidate_phrases(message))

    cleaned_candidates: List[str] = []
    seen_candidates: Set[str] = set()
    for candidate in detail_candidates:
        cleaned = _safe_strip(candidate)
        if cleaned and cleaned.lower() not in seen_candidates:
            seen_candidates.add(cleaned.lower())
            cleaned_candidates.append(cleaned)

    for candidate in cleaned_candidates:
        entry = (
            base_query.filter(AddressBookEntry.name.ilike(f"%{candidate}%"))
            .order_by(AddressBookEntry.name.asc())
            .first()
        )
        if entry:
            return _format_contact_detail(entry)

    keywords = _extract_keywords(message, extra_stop={"contact", "contacts", "address", "book", "list", "show"})
    keyword_conditions = []
    for keyword in keywords[:6]:
        like = f"%{keyword}%"
        keyword_conditions.append(
            or_(
                AddressBookEntry.name.ilike(like),
                AddressBookEntry.company.ilike(like),
                AddressBookEntry.department.ilike(like),
                AddressBookEntry.job_title.ilike(like),
                AddressBookEntry.tags.ilike(like),
                AddressBookEntry.notes.ilike(like),
                AddressBookEntry.city.ilike(like),
            )
        )

    filtered_query = base_query.filter(and_(*keyword_conditions)) if keyword_conditions else base_query
    results = filtered_query.order_by(AddressBookEntry.name.asc()).limit(25).all()
    total = filtered_query.count() if need_total else len(results)

    if not results:
        if filters or keyword_conditions:
            return "No contacts matched those filters."
        return None

    lines = []
    for entry in results:
        company = entry.company or "n/a"
        email = entry.email or "n/a"
        phone = entry.phone or entry.mobile or "n/a"
        location_bits = [bit for bit in [entry.city, entry.country] if bit]
        location = ", ".join(location_bits) if location_bits else "n/a"
        tags = entry.tags or "n/a"
        lines.append(
            f"{entry.name} — {company}; email {email}; phone {phone}; location {location}; tags {tags}"
        )

    header = f"Found {total} contact(s)"
    if filters:
        header += f" ({', '.join(filters)})"
    return header + ":\n" + "\n".join(lines)


def _answer_cross_module_query(message: str, lowered: str, user) -> Optional[str]:
    # Tickets referencing an asset
    if "ticket" in lowered and "asset" in lowered:
        asset = _lookup_hardware_asset_by_identifier(message)
        if not asset:
            for phrase in _extract_candidate_phrases(message):
                fallback = (
                    HardwareAsset.query.filter(
                        or_(
                            HardwareAsset.asset_tag.ilike(f"%{phrase}%"),
                            HardwareAsset.hostname.ilike(f"%{phrase}%"),
                            HardwareAsset.serial_number.ilike(f"%{phrase}%"),
                            HardwareAsset.custom_tag.ilike(f"%{phrase}%"),
                        )
                    )
                    .order_by(HardwareAsset.updated_at.desc())
                    .first()
                )
                if fallback:
                    asset = fallback
                    break
        if asset:
            search_terms = {term for term in [asset.asset_tag, asset.serial_number, asset.hostname, asset.custom_tag] if term}
            if search_terms:
                conditions = []
                for term in list(search_terms)[:5]:
                    like = f"%{term}%"
                    conditions.append(Ticket.subject.ilike(like))
                    conditions.append(Ticket.description.ilike(like))
                ticket_query = Ticket.query.options(joinedload(Ticket.assignee)).filter(or_(*conditions))
                tickets = ticket_query.order_by(Ticket.created_at.desc()).limit(15).all()
                if tickets:
                    lines = []
                    for t in tickets:
                        assignee = t.assignee.username if t.assignee else "Unassigned"
                        status = t.status or "Unknown"
                        lines.append(f"#{t.id} {t.subject} — {status}; assigned to {assignee}")
                    asset_label = asset.asset_tag or asset.hostname or asset.serial_number or asset.model or f"Hardware #{asset.id}"
                    return f"Tickets referencing asset {asset_label}:\n" + "\n".join(lines)
                asset_label = asset.asset_tag or asset.hostname or asset.serial_number or asset.model or f"Hardware #{asset.id}"
                return f"No tickets reference asset {asset_label}."

    # Knowledge articles tied to software
    if "article" in lowered and "software" in lowered:
        software = _lookup_software_asset_by_identifier(message)
        if not software:
            for phrase in _extract_candidate_phrases(message):
                software = SoftwareAsset.query.filter(SoftwareAsset.name.ilike(f"%{phrase}%")).first()
                if software:
                    break
        if software and software.name:
            like = f"%{software.name}%"
            kb_query = KnowledgeArticle.query.filter(
                or_(
                    KnowledgeArticle.title.ilike(like),
                    KnowledgeArticle.summary.ilike(like),
                    KnowledgeArticle.tags.ilike(like),
                    KnowledgeArticle.content.ilike(like),
                )
            ).order_by(KnowledgeArticle.updated_at.desc())
            articles = kb_query.limit(10).all()
            if articles:
                lines = [f"#{article.id} {article.title} — updated {article.updated_at.strftime('%Y-%m-%d') if article.updated_at else 'n/a'}" for article in articles]
                return f"Knowledge articles referencing {software.name}:\n" + "\n".join(lines)
            return f"No knowledge base articles reference {software.name}."

    # Contracts and support for vendor
    if "contract" in lowered and "support" in lowered and "vendor" in lowered:
        vendor_match = VENDOR_PATTERN.search(message)
        if vendor_match:
            vendor_value = _safe_strip(vendor_match.group(1))
            if vendor_value:
                contracts = (
                    Contract.query.filter(func.lower(Contract.vendor) == vendor_value.lower())
                    .order_by(Contract.end_date.asc(), Contract.name.asc())
                    .limit(10)
                    .all()
                )
                if contracts:
                    lines = []
                    for contract in contracts:
                        support = _support_contact_snippet(contract) or "n/a"
                        renewal = contract.renewal_date.strftime("%Y-%m-%d") if contract.renewal_date else "n/a"
                        lines.append(
                            f"{contract.name} — support {support}; renewal {renewal}; auto-renew {'yes' if contract.auto_renew else 'no'}"
                        )
                    return f"Contracts and support for vendor {vendor_value}:\n" + "\n".join(lines)
                return f"No contracts found for vendor {vendor_value}."

    # Assignment lookup by IP
    if "assigned" in lowered and "ip" in lowered:
        ip_match = IP_ADDRESS_PATTERN.search(message)
        if ip_match:
            ip_value = ip_match.group(0)
            lines = []
            host = NetworkHost.query.filter(func.lower(NetworkHost.ip_address) == ip_value.lower()).first()
            if host:
                host_assignee = host.assigned_to or "Unassigned"
                lines.append(
                    f"Network host {ip_value}: hostname {host.hostname or 'n/a'}, assigned to {host_assignee}, reserved {'yes' if host.is_reserved else 'no'}"
                )
            asset = HardwareAsset.query.filter(func.lower(HardwareAsset.ip_address) == ip_value.lower()).first()
            if asset:
                assignee = asset.assignee.username if asset.assignee else "Unassigned"
                label = asset.asset_tag or asset.hostname or asset.serial_number or asset.model or f"Hardware #{asset.id}"
                lines.append(f"Hardware asset {label}: assigned to {assignee}")
            if lines:
                return f"Assignment details for IP {ip_value}:\n" + "\n".join(lines)
            return f"No assignment information found for IP {ip_value}."

    # Combined hardware/software summary for user
    if "hardware" in lowered and "software" in lowered and "user" in lowered:
        target_user = None
        user_match = USER_REF_PATTERN.search(message)
        if user_match:
            target_user = _resolve_user_reference(user_match.group(1))
        elif any(term in lowered for term in ("my hardware", "my software", "for me")) and user:
            target_user = user
        if target_user:
            hardware_assets = HardwareAsset.query.filter(HardwareAsset.assigned_to == target_user.id).all()
            software_assets = SoftwareAsset.query.filter(SoftwareAsset.assigned_to == target_user.id).all()
            lines = [
                f"Hardware assigned: {len(hardware_assets)}",
                f"Software assigned: {len(software_assets)}",
            ]
            if hardware_assets:
                preview = ", ".join(
                    (asset.asset_tag or asset.hostname or asset.serial_number or f"HW#{asset.id}")
                    for asset in hardware_assets[:5]
                )
                if len(hardware_assets) > 5:
                    preview += "…"
                lines.append(f"Hardware preview: {preview}")
            if software_assets:
                preview = ", ".join(
                    (asset.name or asset.custom_tag or f"SW#{asset.id}")
                    for asset in software_assets[:5]
                )
                if len(software_assets) > 5:
                    preview += "…"
                lines.append(f"Software preview: {preview}")
            return f"Assignments for {target_user.username}:\n" + "\n".join(lines)

    return None


def _answer_hardware_query(message: str, lowered: str, user) -> Optional[str]:
    specific = _lookup_hardware_asset_by_identifier(message)
    if specific:
        return _format_hardware_detail(specific)

    base_query = HardwareAsset.query.options(joinedload(HardwareAsset.assignee))
    filters: List[str] = []

    if any(token in lowered for token in ("assigned to me", "my devices", "my hardware", "my laptop")) and user:
        base_query = base_query.filter(HardwareAsset.assigned_to == user.id)
        filters.append(f"assigned to {user.username}")
    else:
        assignee_match = USER_REF_PATTERN.search(message)
        if assignee_match:
            target = _resolve_user_reference(assignee_match.group(1))
            if target:
                base_query = base_query.filter(HardwareAsset.assigned_to == target.id)
                filters.append(f"assigned to {target.username}")

    if "unassigned" in lowered or "available" in lowered:
        base_query = base_query.filter(HardwareAsset.assigned_to.is_(None))
        filters.append("unassigned")

    loc_match = LOCATION_PATTERN.search(message)
    if loc_match:
        location = _safe_strip(loc_match.group(1))
        if location:
            base_query = base_query.filter(func.lower(HardwareAsset.location) == location.lower())
            filters.append(f"location {location}")

    status_match = STATUS_PATTERN.search(message)
    if status_match:
        status_value = _safe_strip(status_match.group(1))
        if status_value:
            base_query = base_query.filter(func.lower(HardwareAsset.status) == status_value.lower())
            filters.append(f"status {status_value}")

    if any(term in lowered for term in ("decommissioned", "retired", "disposed")):
        base_query = base_query.filter(
            or_(
                func.lower(HardwareAsset.status) == "retired",
                func.lower(HardwareAsset.status) == "disposed",
                func.lower(HardwareAsset.status) == "decommissioned",
            )
        )
        filters.append("status decommissioned")

    if "warranty" in lowered or "εγγύηση" in lowered:
        if "out of warranty" in lowered or "εκτός εγγύησης" in lowered:
            base_query = base_query.filter(HardwareAsset.warranty_end.isnot(None))
            base_query = base_query.filter(HardwareAsset.warranty_end < date.today())
            filters.append("out of warranty")
        else:
            date_match = DATE_BY_PATTERN.search(lowered)
            if date_match:
                target_date = _parse_date_string(date_match.group(1))
                if target_date:
                    base_query = base_query.filter(HardwareAsset.warranty_end.isnot(None))
                    base_query = base_query.filter(HardwareAsset.warranty_end <= target_date)
                    filters.append(f"warranty ends by {target_date.isoformat()}")
            elif any(term in lowered for term in ("expiring", "λήγει", "ληγει")):
                window = date.today() + timedelta(days=60)
                base_query = base_query.filter(HardwareAsset.warranty_end.isnot(None))
                base_query = base_query.filter(HardwareAsset.warranty_end <= window)
                filters.append("warranty expiring within 60 days")

    if any(term in lowered for term in ("networked", "have ip", "with ip", "ip address")):
        base_query = base_query.filter(HardwareAsset.ip_address.isnot(None))
        base_query = base_query.filter(HardwareAsset.ip_address != "")
        filters.append("with IP address")

    keywords = _extract_keywords(message, extra_stop={"hardware", "device", "devices", "asset", "assets"})
    need_total = "how many" in lowered or "count" in lowered
    request_all = any(
        phrase in lowered
        for phrase in (
            "all hardware",
            "list hardware",
            "show hardware",
            "hardware inventory",
            "όλο το υλικό",
            "λίστα υλικού",
        )
    ) or lowered.strip() in {"list all hardware", "list hardware", "show all hardware", "display hardware"}

    results: List[HardwareAsset] = []

    if request_all and not keywords and not filters:
        results = base_query.order_by(HardwareAsset.updated_at.desc()).limit(20).all()
        total = base_query.count() if need_total else len(results)
    else:
        phrases = _extract_candidate_phrases(message)
        for phrase in phrases:
            like = f"%{phrase}%"
            phrase_results = (
                base_query.filter(
                    or_(
                        HardwareAsset.category.ilike(like),
                        HardwareAsset.type.ilike(like),
                        HardwareAsset.manufacturer.ilike(like),
                        HardwareAsset.model.ilike(like),
                        HardwareAsset.hostname.ilike(like),
                        HardwareAsset.asset_tag.ilike(like),
                        HardwareAsset.serial_number.ilike(like),
                        HardwareAsset.custom_tag.ilike(like),
                        HardwareAsset.location.ilike(like),
                        HardwareAsset.notes.ilike(like),
                        HardwareAsset.ip_address.ilike(like),
                    )
                )
                .order_by(HardwareAsset.updated_at.desc())
                .limit(10)
                .all()
            )
            if phrase_results:
                results = phrase_results
                break

    if not results:
        keyword_conditions = []
        for keyword in keywords[:6]:
            like = f"%{keyword}%"
            keyword_conditions.append(
                or_(
                    HardwareAsset.category.ilike(like),
                    HardwareAsset.type.ilike(like),
                    HardwareAsset.manufacturer.ilike(like),
                    HardwareAsset.model.ilike(like),
                    HardwareAsset.hostname.ilike(like),
                    HardwareAsset.asset_tag.ilike(like),
                    HardwareAsset.serial_number.ilike(like),
                    HardwareAsset.custom_tag.ilike(like),
                    HardwareAsset.location.ilike(like),
                    HardwareAsset.notes.ilike(like),
                    HardwareAsset.ip_address.ilike(like),
                )
            )

        filtered_query = base_query.filter(or_(*keyword_conditions)) if keyword_conditions else base_query
        results = filtered_query.order_by(HardwareAsset.updated_at.desc()).limit(10).all()
        total = filtered_query.count() if need_total else len(results)
    elif not (request_all and not keywords and not filters):
        total = len(results)

    if not results:
        if request_all and not keywords and not filters:
            fallback_query = HardwareAsset.query.order_by(HardwareAsset.updated_at.desc())
            fallback_results = fallback_query.limit(20).all()
            if fallback_results:
                results = fallback_results
                total = fallback_query.count() if need_total else len(results)

    if not results:
        if filters or keywords or request_all:
            return "No hardware assets matched those filters."
        return None

    lines = []
    for asset in results:
        name = asset.asset_tag or asset.custom_tag or asset.hostname or f"Asset #{asset.id}"
        category = asset.category or asset.type or "hardware"
        status = asset.status or "unknown"
        assignee = asset.assignee.username if asset.assignee else "Unassigned"
        location = asset.location or "n/a"
        line = (
            f"{name} — {category}; status {status}; assigned to {assignee}; location {location}"
        )
        lines.append(line)

    if total > len(results):
        lines.append(f"…and {total - len(results)} more matching assets.")

    if filters:
        header = f"Found {total} hardware asset(s) ({', '.join(filters)})."
    else:
        header = f"Found {total} hardware asset(s)."

    return "\n".join([header] + lines)


def _answer_software_query(message: str, lowered: str, user) -> Optional[str]:
    current_app.logger.info("Assistant software query message=%s", message)
    specific = _lookup_software_asset_by_identifier(message)
    if specific:
        current_app.logger.info("Assistant software query matched specific asset id=%s", specific.id)
        return _format_software_detail(specific)

    if ("perpetual" in lowered and "subscription" in lowered) and ("vs" in lowered or "versus" in lowered):
        total = SoftwareAsset.query.count()
        total_perpetual = SoftwareAsset.query.filter(func.lower(SoftwareAsset.license_type) == "perpetual").count()
        total_subscription = SoftwareAsset.query.filter(func.lower(SoftwareAsset.license_type) == "subscription").count()
        other = max(total - (total_perpetual + total_subscription), 0)
        return (
            f"License types: Perpetual {total_perpetual}, Subscription {total_subscription}, Other {other}, Total {total}."
        )

    base_query = SoftwareAsset.query.options(joinedload(SoftwareAsset.assignee))
    filters: List[str] = []
    keywords = _extract_keywords(message, extra_stop={"software", "license", "licence", "application", "app", "subscription"})
    need_total = "how many" in lowered or "count" in lowered
    request_all = any(
        phrase in lowered
        for phrase in (
            "all software",
            "list software",
            "show software",
            "software inventory",
            "όλο το λογισμικό",
            "λίστα λογισμικού",
        )
    ) or lowered.strip() in {"list all software", "list software", "show all software", "display software"}
    field_filters_applied = False

    category_match = re.search(r"category\s+(?:is|=)?\s*['\"]?([a-z0-9\-\s_/]+)['\"]?", message, flags=re.IGNORECASE)
    if category_match:
        category_value = _safe_strip(category_match.group(1))
        if category_value:
            base_query = base_query.filter(func.lower(SoftwareAsset.category) == category_value.lower())
            filters.append(f"category {category_value}")
            field_filters_applied = True

    vendor_match = re.search(r"vendor\s+(?:is|=)?\s*['\"]?([a-z0-9\-\s_/]+)['\"]?", message, flags=re.IGNORECASE)
    if vendor_match:
        vendor_value = _safe_strip(vendor_match.group(1))
        if vendor_value:
            base_query = base_query.filter(func.lower(SoftwareAsset.vendor) == vendor_value.lower())
            filters.append(f"vendor {vendor_value}")
            field_filters_applied = True

    tag_match = TAG_PATTERN.search(message)
    if tag_match:
        tag_value = _safe_strip(tag_match.group(1))
        if tag_value:
            base_query = base_query.filter(SoftwareAsset.custom_tag.ilike(f"%{tag_value}%"))
            filters.append(f"tag {tag_value}")
            field_filters_applied = True

    if any(token in lowered for token in ("assigned to me", "my software", "my license", "my licences")) and user:
        base_query = base_query.filter(SoftwareAsset.assigned_to == user.id)
        filters.append(f"assigned to {user.username}")
        field_filters_applied = True
    else:
        assignee_match = USER_REF_PATTERN.search(message)
        if assignee_match:
            target = _resolve_user_reference(assignee_match.group(1))
            if target:
                base_query = base_query.filter(SoftwareAsset.assigned_to == target.id)
                filters.append(f"assigned to {target.username}")
                field_filters_applied = True

    if "unassigned" in lowered or "available" in lowered:
        base_query = base_query.filter(SoftwareAsset.assigned_to.is_(None))
        filters.append("unassigned")
        field_filters_applied = True

    expiration_date_match = None
    if any(term in lowered for term in ("expir", "expires", "λήγει", "ληγει")):
        expiration_date_match = DATE_BY_PATTERN.search(lowered)

    if expiration_date_match:
        target_date = _parse_date_string(expiration_date_match.group(1))
        if target_date:
            base_query = base_query.filter(SoftwareAsset.expiration_date.isnot(None))
            base_query = base_query.filter(SoftwareAsset.expiration_date <= target_date)
            filters.append(f"expires by {target_date.isoformat()}")
            field_filters_applied = True
    elif any(term in lowered for term in ("expir", "expires", "λήγει", "ληγει")):
        today = date.today()
        window = today + timedelta(days=60)
        base_query = base_query.filter(
            SoftwareAsset.expiration_date.isnot(None),
            SoftwareAsset.expiration_date <= window,
        )
        filters.append("expiring within 60 days")
        field_filters_applied = True

    results: List[SoftwareAsset] = []
    if request_all and not field_filters_applied:
        results = base_query.order_by(SoftwareAsset.updated_at.desc()).limit(20).all()
        total = base_query.count() if need_total else len(results)
        current_app.logger.info("Assistant software query request_all fetched=%s total=%s", len(results), total)
    else:
        phrases = _extract_candidate_phrases(message)
        for phrase in phrases:
            like = f"%{phrase}%"
            name_results = (
                base_query.filter(
                    or_(
                        SoftwareAsset.name.ilike(like),
                        SoftwareAsset.custom_tag.ilike(like),
                    )
                )
                .order_by(SoftwareAsset.updated_at.desc())
                .limit(10)
                .all()
            )
            if name_results:
                results = name_results
                current_app.logger.info("Assistant software query phrase matched by name/custom_tag phrase=%s count=%s", phrase, len(results))
                break

            phrase_results = (
                base_query.filter(
                    or_(*[column.ilike(like) for column in SOFTWARE_SEARCH_COLUMNS])
                )
                .order_by(SoftwareAsset.updated_at.desc())
                .limit(10)
                .all()
            )
            if phrase_results:
                results = phrase_results
                current_app.logger.info("Assistant software query phrase matched by broad search phrase=%s count=%s", phrase, len(results))
                break

    if not results:
        keyword_conditions = []
        for keyword in keywords[:6]:
            like = f"%{keyword}%"
            keyword_conditions.append(or_(*[column.ilike(like) for column in SOFTWARE_SEARCH_COLUMNS]))

        filtered_query = base_query.filter(and_(*keyword_conditions)) if keyword_conditions else base_query
        results = filtered_query.order_by(SoftwareAsset.updated_at.desc()).limit(10).all()
        total = filtered_query.count() if need_total else len(results)
    elif not (request_all and not field_filters_applied):
        total = len(results)

    if not results:
        if field_filters_applied:
            fallback_query = base_query.order_by(SoftwareAsset.updated_at.desc())
            fallback_results = fallback_query.limit(20).all()
            if fallback_results:
                current_app.logger.info("Assistant software query fallback (field filters) returned %s records", len(fallback_results))
                results = fallback_results
                total = fallback_query.count() if need_total else len(results)
        elif request_all or (not field_filters_applied and not keywords):
            fallback_query = SoftwareAsset.query.order_by(SoftwareAsset.updated_at.desc())
            fallback_results = fallback_query.limit(20 if request_all else 10).all()
            if fallback_results:
                current_app.logger.info("Assistant software query fallback (global) returned %s records", len(fallback_results))
                results = fallback_results
                total = fallback_query.count() if need_total else len(results)

    if not results:
        current_app.logger.info("Assistant software query empty after fallbacks.")
        if filters or keywords or request_all:
            return "No software assets matched those filters."
        return None

    lines = []
    for asset in results:
        name = asset.name or f"Software #{asset.id}"
        version = asset.version or "—"
        vendor = asset.vendor or "Unknown vendor"
        assignee = asset.assignee.username if asset.assignee else "Unassigned"
        expires = asset.expiration_date.strftime("%Y-%m-%d") if asset.expiration_date else "n/a"
        license_key = asset.license_key or "n/a"
        serial = asset.serial_number or "n/a"
        line = (
            f"{name} {version} — {vendor}; license type {asset.license_type or 'n/a'}; license key {license_key}; serial {serial}; assigned to {assignee}; expires {expires}"
        )
        lines.append(line)

    if total > len(results):
        lines.append(f"…and {total - len(results)} more matching licenses.")

    if filters:
        header = f"Found {total} software record(s) ({', '.join(filters)})."
    else:
        header = f"Found {total} software record(s)."

    current_app.logger.info("Assistant software query returning %s rows", len(results))
    return "\n".join([header] + lines)


def _extract_keywords(text: str, extra_stop: Optional[Iterable[str]] = None) -> List[str]:
    stops = set(STOP_WORDS)
    if extra_stop:
        stops.update(extra_stop)
    words = re.findall(r"[a-z0-9]{3,}", text.lower())
    return [word for word in words if word not in stops]


DATE_PARSE_FORMATS = ("%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y")


def _parse_date_string(raw: Optional[str]) -> Optional[date]:
    if not raw:
        return None
    cleaned = _safe_strip(raw)
    if not cleaned:
        return None
    for fmt in DATE_PARSE_FORMATS:
        try:
            return datetime.strptime(cleaned, fmt).date()
        except ValueError:
            continue
    return None


def _parse_decimal(raw: Optional[str]) -> Optional[Decimal]:
    if not raw:
        return None
    cleaned = _safe_strip(raw)
    if not cleaned:
        return None
    normalized = cleaned.replace(" ", "").replace(",", "")
    try:
        return Decimal(normalized)
    except (InvalidOperation, ValueError):
        return None


def _support_contact_snippet(contract: Contract) -> str:
    parts: List[str] = []
    if contract.support_email:
        parts.append(f"email {contract.support_email}")
    if contract.support_phone:
        parts.append(f"phone {contract.support_phone}")
    if contract.support_url:
        parts.append(f"url {contract.support_url}")
    return ", ".join(parts)


def _format_contract_detail(contract: Contract, include_support: bool = False) -> str:
    owner_name = contract.owner.username if contract.owner else "n/a"
    start_date = contract.start_date.strftime("%Y-%m-%d") if contract.start_date else "n/a"
    end_date = contract.end_date.strftime("%Y-%m-%d") if contract.end_date else "n/a"
    renewal = contract.renewal_date.strftime("%Y-%m-%d") if contract.renewal_date else "n/a"
    value = f"{contract.value} {contract.currency}" if contract.value else "n/a"
    lines = [
        f"Contract #{contract.id}: {contract.name}",
        f"Type: {contract.contract_type} | Status: {contract.status or 'n/a'} | Vendor: {contract.vendor or 'n/a'}",
        f"Contract number: {contract.contract_number or 'n/a'} | PO: {contract.po_number or 'n/a'} | Owner: {owner_name}",
        f"Start: {start_date} | End: {end_date} | Renewal: {renewal} | Auto-renew: {'yes' if contract.auto_renew else 'no'}",
        f"Value: {value}",
    ]
    support_info = _support_contact_snippet(contract)
    if include_support or support_info:
        lines.append(f"Support: {support_info or 'n/a'}")
    if contract.coverage_scope:
        lines.append(f"Coverage: {contract.coverage_scope}")
    if contract.notes:
        lines.append(f"Notes: {contract.notes}")
    return "\n".join(lines)


def _format_contact_detail(entry: AddressBookEntry) -> str:
    lines = [
        f"Contact #{entry.id}: {entry.name}",
        f"Company: {entry.company or 'n/a'} | Department: {entry.department or 'n/a'} | Job title: {entry.job_title or 'n/a'}",
        f"Email: {entry.email or 'n/a'} | Phone: {entry.phone or 'n/a'} | Mobile: {entry.mobile or 'n/a'}",
        f"Category: {entry.category or 'n/a'} | City: {entry.city or 'n/a'} | Country: {entry.country or 'n/a'}",
    ]
    if entry.tags:
        lines.append(f"Tags: {entry.tags}")
    if entry.notes:
        lines.append(f"Notes: {entry.notes}")
    if entry.website:
        lines.append(f"Website: {entry.website}")
    if entry.address_line or entry.postal_code or entry.state:
        address_bits = [bit for bit in [entry.address_line, entry.state, entry.postal_code] if bit]
        if address_bits:
            lines.append("Address: " + ", ".join(address_bits))
    return "\n".join(lines)


def _extract_field_terms(message: str, *patterns: re.Pattern) -> List[str]:
    terms: List[str] = []
    for pattern in patterns:
        if not pattern:
            continue
        for match in pattern.finditer(message):
            term = _safe_strip(match.group(1))
            if term:
                terms.append(term)
    deduped: List[str] = []
    seen: Set[str] = set()
    for term in terms:
        key = term.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(term)
    return deduped


def _clean_user_identifier(identifier: str) -> str:
    cleaned = _safe_strip(identifier, " \"'.,:;!?")
    if not cleaned:
        return ""
    lowered = cleaned.lower()
    for delimiter in USER_REFERENCE_DELIMITERS:
        idx = lowered.find(delimiter)
        if idx != -1:
            cleaned = cleaned[:idx]
            lowered = cleaned.lower()
    cleaned = re.split(r"\s+(?:tickets?|incidents?|requests?)\b", cleaned, maxsplit=1)[0]
    return _safe_strip(cleaned)


def _resolve_user_reference(identifier: str) -> Optional[User]:
    cleaned = _clean_user_identifier(identifier)
    if not cleaned:
        return None
    return User.query.filter(func.lower(User.username) == cleaned.lower()).first()


def _extract_candidate_phrases(message: str) -> List[str]:
    candidates: List[str] = []

    # quoted phrases take priority
    for match in QUOTED_PATTERN.finditer(message):
        phrase = _safe_strip(match.group(1) or match.group(2))
        if phrase:
            candidates.append(phrase)

    # phrases introduced by keywords (for, find, etc.)
    for match in PHRASE_PATTERN.finditer(message):
        raw = match.group(1) or ""
        phrase = _safe_strip(raw)
        if not phrase:
            continue
        phrase = re.split(r"[?.,;]", phrase)[0]
        phrase = re.sub(r"\b(key|keys|license|licence|serial|product key)\b", "", phrase, flags=re.IGNORECASE)
        phrase = _safe_strip(re.sub(r"\s+", " ", phrase))
        if phrase and phrase.lower() not in (cand.lower() for cand in candidates):
            candidates.append(phrase)

    return candidates[:3]


def _extract_user_tokens(message: str) -> List[str]:
    tokens: List[str] = []

    for match in USER_REF_PATTERN.finditer(message):
        candidate = _safe_strip(match.group(1), " \"'.,:;!")
        if candidate:
            tokens.append(candidate)

    for match in QUOTED_PATTERN.finditer(message):
        candidate = _safe_strip(match.group(1) or match.group(2))
        if candidate:
            tokens.append(candidate)

    seen = set()
    deduped: List[str] = []
    for token in tokens:
        key = token.lower()
        if key and key not in seen:
            seen.add(key)
            deduped.append(token)
    return deduped


def _extract_network_host_tokens(message: str) -> List[str]:
    tokens: List[str] = []

    for match in HOSTNAME_QUERY_PATTERN.finditer(message):
        candidate = _safe_strip(match.group(1))
        if candidate:
            tokens.append(candidate)

    for match in MAC_QUERY_PATTERN.finditer(message):
        candidate = _safe_strip(match.group(1))
        if candidate:
            tokens.append(candidate)

    for match in IP_ADDRESS_PATTERN.finditer(message):
        candidate = _safe_strip(match.group(0))
        if candidate and "/" not in candidate:  # avoid double-counting CIDRs
            tokens.append(candidate)

    return tokens


def _normalize_network_name_hint(raw: str) -> str:
    text = _safe_strip(raw)
    if not text:
        return ""
    breaker = NETWORK_HINT_BREAK_RE.search(text)

    def _tokenize(segment: str) -> List[str]:
        cleaned = re.sub(r"[^\w\s\-.:/]", " ", segment)
        return [word for word in re.split(r"\s+", cleaned) if word and word.lower() not in NETWORK_HINT_STOPWORDS]

    if breaker:
        before = text[: breaker.start()]
        after = text[breaker.end() :]
        after_tokens = _tokenize(after)
        before_tokens = _tokenize(before)
        if after_tokens:
            words = after_tokens
        elif before_tokens:
            words = before_tokens
        else:
            words = _tokenize(text)
    else:
        words = _tokenize(text)

    if not words:
        return ""
    cleaned = re.sub(r"\s+", " ", " ".join(words[:6])).strip(" -_/")
    return cleaned


def _canonicalize_cidr(value: str) -> str:
    try:
        return str(ipaddress.ip_network(value, strict=False))
    except ValueError:
        return (value or "").strip().lower()


def _find_network_by_cidr(networks: Iterable[Network], cidr_value: str) -> Optional[Network]:
    canonical = _canonicalize_cidr(cidr_value)
    for net in networks:
        if _canonicalize_cidr(net.cidr or "") == canonical:
            return net
    for net in networks:
        if canonical in (net.cidr or "").lower():
            return net
    return None


def _normalize_label(value: str) -> str:
    if not value:
        return ""
    cleaned = re.sub(r"[^a-z0-9]+", " ", value.lower())
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _match_network_by_normalized(networks: Iterable[Network], normalized: str) -> Optional[Network]:
    if not normalized:
        return None

    lowered = normalized.lower().strip()
    variations = {lowered}
    if not lowered.endswith(" network"):
        variations.add(f"{lowered} network")
    if not lowered.endswith(" subnet"):
        variations.add(f"{lowered} subnet")
    if not lowered.endswith(" vlan"):
        variations.add(f"{lowered} vlan")

    for value in variations:
        for net in networks:
            if _normalize_label(net.name) == _normalize_label(value):
                return net
            if (net.name or "").strip().lower() == value:
                return net

    tokens = [token for token in re.split(r"\s+", normalized) if token]
    if tokens:
        best_match: Optional[Tuple[int, Network]] = None
        for net in networks:
            name_norm = _normalize_label(net.name)
            token_hits = sum(1 for token in tokens if token in name_norm)
            if token_hits <= 0:
                continue
            if not best_match or token_hits > best_match[0]:
                best_match = (token_hits, net)
        if best_match:
            return best_match[1]

    lowered_norm = _normalize_label(normalized)
    for net in networks:
        for field in (net.site, net.vlan, net.description, net.notes):
            if lowered_norm and lowered_norm in _normalize_label(field or ""):
                return net

    for net in networks:
        name_lower = (net.name or "").lower()
        if lowered in name_lower:
            return net

    return None


def _lookup_hardware_asset_by_identifier(message: str) -> Optional[HardwareAsset]:
    query = HardwareAsset.query.options(joinedload(HardwareAsset.assignee))

    for pattern, column in (
        (ASSET_TAG_PATTERN, HardwareAsset.asset_tag),
        (HOSTNAME_PATTERN, HardwareAsset.hostname),
        (SERIAL_PATTERN, HardwareAsset.serial_number),
    ):
        match = pattern.search(message)
        if match:
            value = _safe_strip(match.group(1))
            if value:
                asset = query.filter(func.lower(column) == value.lower()).first()
                if asset:
                    return asset

    ip_match = IP_ADDRESS_PATTERN.search(message)
    if ip_match:
        value = _safe_strip(ip_match.group(0))
        if value:
            asset = query.filter(func.lower(HardwareAsset.ip_address) == value.lower()).first()
            if asset:
                return asset

    return None


def _lookup_software_asset_by_identifier(message: str) -> Optional[SoftwareAsset]:
    query = SoftwareAsset.query.options(joinedload(SoftwareAsset.assignee))

    candidates: List[Tuple[str, Any]] = []

    for match in SOFTWARE_TAG_PATTERN.finditer(message):
        raw = match.group(1)
        value = _safe_strip(raw)
        if value:
            candidates.append((value, SoftwareAsset.custom_tag))

    direct_patterns = (
        (
            re.compile(
                r"license\s+(?:key|code|number|id)\s*[:#]?\s*([a-z0-9\-]{4,})",
                re.IGNORECASE,
            ),
            SoftwareAsset.license_key,
        ),
        (
            re.compile(
                r"activation\s+(?:code|key)\s*[:#]?\s*([a-z0-9\-]{4,})",
                re.IGNORECASE,
            ),
            SoftwareAsset.license_key,
        ),
        (
            re.compile(
                r"seria\w*\s+(?:number|no\.?)\s*[:#]?\s*([a-z0-9_.\-]+)",
                re.IGNORECASE,
            ),
            SoftwareAsset.serial_number,
        ),
        (
            re.compile(
                r"serial\s*#\s*([a-z0-9_.\-]+)",
                re.IGNORECASE,
            ),
            SoftwareAsset.serial_number,
        ),
    )

    for pattern, column in direct_patterns:
        for match in pattern.finditer(message):
            raw = match.group(1)
            cleaned = _safe_strip(raw)
            value = cleaned.strip(" '\"#:;,") if cleaned else ""
            if value:
                candidates.append((value, column))

    bare_license_pattern = re.compile(r"\b(?:[a-z0-9]{4,}-){2,}[a-z0-9]{4,}\b", re.IGNORECASE)
    for match in bare_license_pattern.finditer(message):
        raw = match.group(0)
        cleaned = _safe_strip(raw)
        value = cleaned.strip(" '\"#:;,") if cleaned else ""
        if value:
            candidates.append((value, SoftwareAsset.license_key))

    seen: set[str] = set()
    for value, column in candidates:
        lowered = value.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        asset = query.filter(func.lower(column) == lowered).first()
        if asset:
            return asset

    return None


def _format_hardware_detail(asset: HardwareAsset) -> str:
    assignee = asset.assignee.username if asset.assignee else "Unassigned"
    created = asset.created_at.strftime("%Y-%m-%d") if asset.created_at else "—"
    updated = asset.updated_at.strftime("%Y-%m-%d") if asset.updated_at else "—"
    return (
        f"Hardware asset {asset.asset_tag or asset.custom_tag or asset.hostname or '#' + str(asset.id)}\n"
        f"Category: {asset.category or asset.type or 'n/a'} | Manufacturer: {asset.manufacturer or 'n/a'} | Model: {asset.model or 'n/a'}\n"
        f"Assigned to: {assignee} | Location: {asset.location or 'n/a'}\n"
        f"Status: {asset.status or 'n/a'} | Created: {created} | Updated: {updated}"
    )


def _format_software_detail(asset: SoftwareAsset) -> str:
    assignee = asset.assignee.username if asset.assignee else "Unassigned"
    expires = asset.expiration_date.strftime("%Y-%m-%d") if asset.expiration_date else "n/a"
    renewed = asset.renewal_date.strftime("%Y-%m-%d") if asset.renewal_date else "n/a"
    license_key = asset.license_key or "n/a"
    return (
        f"Software asset {asset.name or asset.custom_tag or '#' + str(asset.id)}\n"
        f"Version: {asset.version or 'n/a'} | Vendor: {asset.vendor or 'n/a'} | License type: {asset.license_type or 'n/a'}\n"
        f"License key: {license_key}\n"
        f"Assigned to: {assignee} | Expires: {expires} | Renewal date: {renewed}"
    )
